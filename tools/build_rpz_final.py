from __future__ import annotations

import math
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "rpz"
IMG = OUT / "images"
DOCX = OUT / "RPZ_JobAbility.docx"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


F_TITLE = font(44, True)
F_H = font(28, True)
F_B = font(22, True)
F = font(20)
F_S = font(16)
F_XS = font(14)


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        line = ""
        for word in raw.split():
            candidate = (line + " " + word).strip()
            if draw.textlength(candidate, font=fnt) <= width:
                line = candidate
            else:
                if line:
                    lines.append(line)
                line = word
        if line:
            lines.append(line)
    return lines or [""]


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, body: str, color: str) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle((x1 + 8, y1 + 8, x2 + 8, y2 + 8), 12, fill="#D7DEE8")
    draw.rounded_rectangle(xy, 12, fill="#FFFFFF", outline=color, width=3)
    draw.rectangle((x1, y1, x2, y1 + 42), fill=color)
    draw.text((x1 + 18, y1 + 10), title, font=F_B, fill="#FFFFFF")
    yy = y1 + 60
    for line in wrap(draw, body, F_S, x2 - x1 - 36):
        draw.text((x1 + 18, yy), line, font=F_S, fill="#111827")
        yy += 22


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    pts: list[tuple[int, int]],
    label: str = "",
    color: str = "#334155",
    label_pos: tuple[int, int] | None = None,
) -> None:
    draw.line(pts, fill=color, width=4, joint="curve")
    x1, y1 = pts[-2]
    x2, y2 = pts[-1]
    ang = math.atan2(y2 - y1, x2 - x1)
    size = 16
    p1 = (x2 - size * math.cos(ang - 0.45), y2 - size * math.sin(ang - 0.45))
    p2 = (x2 - size * math.cos(ang + 0.45), y2 - size * math.sin(ang + 0.45))
    draw.polygon([(x2, y2), p1, p2], fill=color)
    if label:
        if label_pos is None:
            label_pos = pts[len(pts) // 2]
        tx, ty = label_pos
        pad = 5
        bbox = draw.textbbox((tx, ty), label, font=F_XS, anchor="mm")
        draw.rounded_rectangle((bbox[0] - pad, bbox[1] - pad, bbox[2] + pad, bbox[3] + pad), 7, fill="#FFFFFF", outline="#CBD5E1")
        draw.text((tx, ty), label, font=F_XS, fill="#334155", anchor="mm")


def make_er(path: Path) -> None:
    img = Image.new("RGB", (2100, 1250), "#F7F8FA")
    draw = ImageDraw.Draw(img)
    draw.text((1050, 45), "JobAbility - ER-диаграмма основных таблиц", font=F_TITLE, fill="#0F172A", anchor="ma")

    def entity(x: int, y: int, w: int, title: str, fields: list[tuple[str, str]]) -> tuple[int, int, int, int]:
        row_h = 38
        h = 52 + row_h * len(fields) + 16
        draw.rounded_rectangle((x + 8, y + 8, x + w + 8, y + h + 8), 10, fill="#E5E7EB")
        draw.rounded_rectangle((x, y, x + w, y + h), 10, fill="#F3F4F6", outline="#D1D5DB", width=2)
        draw.rounded_rectangle((x, y, x + w, y + 48), 10, fill="#176496", outline="#176496")
        draw.rectangle((x, y + 32, x + w, y + 48), fill="#176496")
        draw.text((x + 16, y + 12), title, font=F_B, fill="#FFFFFF")
        yy = y + 64
        for name, typ in fields:
            draw.text((x + 18, yy), name + ("  key" if name == "id" else ""), font=F_S, fill="#111827")
            draw.text((x + w - 18, yy), typ, font=F_S, fill="#4B5563", anchor="ra")
            yy += row_h
        return x, y, x + w, y + h

    user = entity(80, 790, 330, "auth_user", [("id", "int"), ("username", "varchar"), ("email", "varchar")])
    profile = entity(560, 125, 455, "core_applicantprofile", [
        ("id", "bigint"), ("full_name", "varchar(255)"), ("phone", "varchar(64)"), ("city", "varchar(255)"),
        ("age", "int"), ("gender", "varchar(16)"), ("disability_category", "varchar(16)"), ("user_id", "int NN")
    ])
    app = entity(1110, 135, 430, "core_application", [
        ("id", "bigint"), ("status", "varchar(16)"), ("created_at", "timestamptz"), ("formed_at", "timestamptz"),
        ("completed_at", "timestamptz"), ("total_salary", "int"), ("applicant_id", "bigint"), ("creator_id", "int NN")
    ])
    line = entity(1645, 150, 370, "core_applicationvacancy", [
        ("id", "bigint"), ("qty", "int"), ("comment", "text"), ("is_main", "boolean"),
        ("order_index", "int"), ("application_id", "bigint NN"), ("vacancy_id", "bigint NN")
    ])
    vacancy = entity(760, 735, 505, "core_vacancy", [
        ("id", "bigint"), ("title", "varchar(255)"), ("company", "varchar(255)"), ("city", "varchar(255)"),
        ("salary", "int"), ("description", "text"), ("is_active", "boolean"), ("image", "varchar(100)"),
        ("disability_support", "varchar(255)"), ("schedule", "varchar(255)")
    ])

    draw_arrow(draw, [(user[2], 860), (500, 860), (500, 310), (profile[0], 310)], "1 : 0..1", "#9CA3AF", (420, 710))
    draw_arrow(draw, [(user[2], 900), (640, 900), (640, 1010), (vacancy[0], 1010)], "1 : N", "#9CA3AF", (600, 950))
    draw_arrow(draw, [(profile[2], 300), (app[0], 300)], "1 : N", "#9CA3AF", (1060, 270))
    draw_arrow(draw, [(app[2], 325), (line[0], 325)], "1 : N", "#9CA3AF", (1588, 295))
    draw_arrow(draw, [(line[0] + 185, line[3]), (line[0] + 185, 1000), (vacancy[2], 1000)], "N : 1", "#9CA3AF", (1520, 960))
    draw_arrow(draw, [(vacancy[0], 875), (500, 875), (500, 900), (user[2], 900)], "creator", "#9CA3AF", (565, 835))
    img.save(path)


def make_bpmn(path: Path) -> None:
    img = Image.new("RGB", (4200, 2200), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    f_lane = font(28, True)
    f_task = font(24)
    f_note = font(18)
    draw.text((2100, 42), "JobAbility - детализированный BPMN процесса трудоустройства", font=F_TITLE, anchor="ma", fill="#0F172A")
    lanes = [
        ("Соискатель", 140, 680, "#F8FAFC"),
        ("Работодатель", 680, 1060, "#F8FAFC"),
        ("Модератор", 1060, 1450, "#F8FAFC"),
        ("Сервис JobAbility", 1450, 2020, "#F8FAFC"),
    ]
    for title, y1, y2, fill in lanes:
        draw.rectangle((100, y1, 4100, y2), fill=fill, outline="#111827", width=2)
        draw.rectangle((100, y1, 175, y2), fill="#FFFFFF", outline="#111827", width=2)
        draw.text((137, (y1 + y2) // 2), title, font=f_lane, anchor="mm", fill="#0F172A")

    def task(name: str, x: int, y: int, text: str, w: int = 320, h: int = 112):
        draw.rounded_rectangle((x - w // 2, y - h // 2, x + w // 2, y + h // 2), 18, fill="#FFFFFF", outline="#111827", width=3)
        yy = y - 30
        for line in wrap(draw, text, f_task, w - 34):
            draw.text((x, yy), line, font=f_task, fill="#111827", anchor="ma")
            yy += 27
        return name, (x, y)

    def event(name: str, x: int, y: int, text: str, end: bool = False):
        draw.ellipse((x - 36, y - 36, x + 36, y + 36), fill="#FFFFFF", outline="#111827", width=5 if end else 3)
        draw.text((x, y + 58), text, font=f_note, anchor="ma", fill="#111827")
        return name, (x, y)

    def gateway(name: str, x: int, y: int, text: str):
        draw.polygon([(x, y - 48), (x + 48, y), (x, y + 48), (x - 48, y)], fill="#FFFFFF", outline="#111827")
        draw.text((x, y + 74), text, font=f_note, anchor="ma", fill="#111827")
        return name, (x, y)

    nodes = dict([
        event("start", 280, 400, "Старт"),
        task("register", 560, 400, "Регистрация или вход"),
        task("catalog", 890, 400, "Просмотр списка вакансий"),
        task("detail", 1220, 400, "Открытие карточки вакансии"),
        task("cart", 1560, 400, "Добавление вакансии в черновик"),
        task("profile", 1900, 400, "Заполнение профиля и заявки"),
        gateway("ready", 2240, 400, "Данные полные?"),
        task("form", 2580, 400, "Формирование заявки"),
        task("track", 2920, 400, "Просмотр статуса заявки"),
        event("finishUser", 3220, 400, "Уведомление", True),
        task("createVac", 620, 870, "Создание вакансии"),
        task("waitModeration", 980, 870, "Ожидание модерации"),
        task("responses", 2920, 870, "Просмотр откликов"),
        task("modLogin", 620, 1245, "Вход модератора"),
        task("modVac", 980, 1245, "Проверка вакансии"),
        gateway("vacOk", 1320, 1245, "Вакансия корректна?"),
        task("modApps", 2240, 1245, "Список сформированных заявок"),
        task("appDecision", 2580, 1245, "Завершение или отклонение заявки"),
        task("auth", 560, 1660, "AuthDomain: сессия, роли, профиль"),
        task("vacDomain", 1030, 1660, "VacancyDomain: каталог, фильтры, модерация"),
        task("appDomain", 1510, 1660, "ApplicationDomain: черновик, строки, расчет"),
        task("notify", 1990, 1660, "SystemDomain: Swagger, metrics, audit"),
        task("storage", 2500, 1660, "PostgreSQL + MinIO + Redis"),
    ])

    flow = [
        ("start", "register", "POST /api/users/login|register"),
        ("register", "catalog", "GET /api/vacancies"),
        ("catalog", "detail", "GET /api/vacancies/{id}"),
        ("detail", "cart", "POST /api/application-lines"),
        ("cart", "profile", "GET/PUT /api/users/profile"),
        ("profile", "ready", ""),
        ("ready", "form", "да: PUT /api/applications/{id}/form"),
        ("form", "track", "GET /api/applications"),
        ("track", "finishUser", "статус FINISHED/REJECTED"),
        ("createVac", "waitModeration", "POST /api/vacancies"),
        ("modLogin", "modVac", "GET /api/vacancies/pending"),
        ("modVac", "vacOk", ""),
        ("vacOk", "catalog", "approve: PUT /moderate"),
        ("vacOk", "createVac", "reject: исправить"),
        ("modApps", "appDecision", "PUT /api/applications/{id}/moderate"),
        ("appDecision", "responses", "GET /api/applications/employer-responses"),
        ("register", "auth", "session"),
        ("catalog", "vacDomain", "filters"),
        ("cart", "appDomain", "draft"),
        ("form", "appDomain", "total_salary"),
        ("appDecision", "notify", "audit/metrics"),
        ("appDomain", "storage", "SQL/S3/cache"),
        ("vacDomain", "storage", "SQL/media"),
    ]
    for a, b, label in flow:
        x1, y1 = nodes[a]
        x2, y2 = nodes[b]
        mid_y = y1
        if abs(y2 - y1) > 190:
            mid_y = (y1 + y2) // 2
            pts = [(x1 + 140, y1), (x1 + 190, y1), (x1 + 190, mid_y), (x2 - 160, mid_y), (x2 - 160, y2), (x2 - 140, y2)]
        else:
            pts = [(x1 + 140, y1), (x2 - 140, y2)]
        draw_arrow(draw, pts, label, "#334155")
    img.save(path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(section) -> None:
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p._p.clear_content()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld)
    run._r.append(instr)
    run._r.append(fld2)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_page_number(section)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    for name, size in [("Heading 1", 16), ("Heading 2", 15), ("Heading 3", 14)]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(15, 23, 42)
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(6)


def p(doc: Document, text: str = "", bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.first_line_indent = Cm(1.25) if align == WD_ALIGN_PARAGRAPH.JUSTIFY else Cm(0)
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = par.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    run.font.bold = bold


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def formula(doc: Document, expr: str, number: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(13.5)
    table.columns[1].width = Cm(2.0)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 80, 80, 80, 80)
        tc_pr = cell._tc.get_or_add_tcPr()
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tc_pr.append(border)
    left, right = table.rows[0].cells
    par = left.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.first_line_indent = Cm(0)
    r = par.add_run(expr)
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    r.font.size = Pt(15)
    par2 = right.paragraphs[0]
    par2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    par2.paragraph_format.first_line_indent = Cm(0)
    r2 = par2.add_run(number)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(14)


def table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size: int = 10) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(4)
    rr = cap.add_run(caption)
    rr.font.name = "Times New Roman"
    rr.font.size = Pt(12)
    rr.font.bold = True
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    if widths:
        for i, w in enumerate(widths):
            tbl.columns[i].width = Cm(w)
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_shading(cell, "EAF2F8")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        par = cell.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            par = cell.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(str(value)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.first_line_indent = Cm(0)
            run = par.add_run(str(value))
            run.font.name = "Times New Roman"
            run.font.size = Pt(font_size)


def picture(doc: Document, path: Path, caption: str, width_cm: float = 15.5) -> None:
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.first_line_indent = Cm(0)
    par.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    r = cap.add_run(caption)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.italic = True


def add_landscape(doc: Document):
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)
    section.footer_distance = Cm(1.0)
    add_page_number(section)


def add_portrait(doc: Document):
    section = doc.add_section()
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    add_page_number(section)


def build() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    IMG.mkdir(parents=True, exist_ok=True)
    er = IMG / "er_final.png"
    bpmn = IMG / "bpmn_final.png"
    make_er(er)
    make_bpmn(bpmn)
    use_case = IMG / "use_case.png"
    state = IMG / "state.png"
    class_img = ROOT / "docs" / "diagram_images" / "jobability_class_diagram.png"
    deploy_img = ROOT / "docs" / "diagram_images" / "jobability_deployment_diagram.png"
    seq_img = ROOT / "docs" / "diagram_images" / "jobability_sequence_diagram.png"
    seq_part_1 = ROOT / "docs" / "diagram_images" / "jobability_sequence_diagram_part_1.png"
    seq_part_2 = ROOT / "docs" / "diagram_images" / "jobability_sequence_diagram_part_2.png"

    doc = Document()
    configure(doc)

    for line in [
        "Министерство науки и высшего образования Российской Федерации",
        "МГТУ им. Н. Э. Баумана",
        "Кафедра ИУ-5 «Системы обработки информации и управления»",
    ]:
        p(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    r = title.add_run("РАСЧЕТНО-ПОЯСНИТЕЛЬНАЯ ЗАПИСКА\nк итоговому проекту\n«JobAbility»")
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r.font.bold = True
    p(doc, "Система трудоустройства для людей с ограниченными возможностями", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    for line in ["Студент: Илья", "Группа: ________", "Преподаватель: ________", "Москва, 2026"]:
        p(doc, line, align=WD_ALIGN_PARAGRAPH.RIGHT if line != "Москва, 2026" else WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    heading(doc, "Введение")
    intro = [
        "Актуальность проекта JobAbility связана с тем, что цифровой рынок труда должен учитывать не только поиск вакансий по профессии и зарплате, но и реальные условия доступности рабочего места. По данным Всемирной организации здравоохранения, около 1,3 млрд человек, то есть примерно 16 % населения мира, живут со значительными ограничениями здоровья. Поэтому задача инклюзивного трудоустройства является массовой социальной и технологической задачей, а не узкой вспомогательной функцией кадрового сервиса [1].",
        "В российском правовом поле занятость инвалидов регулируется Федеральным законом от 24.11.1995 № 181-ФЗ, в котором закреплены меры социальной защиты, квотирование рабочих мест и необходимость специальных условий труда [3]. На практике это означает, что работодатель должен не просто разместить вакансию, а описать график, формат работы, доступность офиса, возможность удаленной занятости, требования к коммуникации и поддержку ассистивных технологий.",
        "Цель работы — разработать и развернуть веб-систему JobAbility для публикации доступных вакансий, регистрации соискателей и работодателей, формирования заявок, модерации вакансий и заявок, хранения медиафайлов и автоматического обновления приложения через CI/CD в Kubernetes.",
        "Назначение системы состоит в том, чтобы объединить в одном информационном контуре трех участников: соискателя, которому нужен понятный каталог и безопасная заявка; работодателя, которому нужен инструмент публикации вакансий; модератора, который контролирует качество данных и завершает обработку заявок.",
        "К нефункциональным требованиям относятся: адаптивность интерфейса на desktop и mobile; ролевая модель доступа; хранение структурированных данных в PostgreSQL, медиафайлов в MinIO и временных данных в Redis; воспроизводимое развертывание в Docker/k3s; автоматическая сборка, тестирование и доставка через GitLab CI/CD; наблюдаемость через Prometheus и Grafana; приемлемое время ответа при демонстрационной нагрузке до 10 RPS.",
        "Задачи работы включали анализ предметной области, проектирование бизнес-процессов и диаграмм, реализацию backend на Django REST Framework, frontend на React, Dockerfile для backend и frontend, Kubernetes-манифесты для всех сервисов, настройку GitLab Runner, подготовку pipeline build-test-upload-deploy, заполнение демонстрационных данных и оформление РПЗ.",
    ]
    for text in intro:
        p(doc, text)

    heading(doc, "Предметная область")
    domain = [
        "Предметная область рассматривается на примере условной компании «JobAbility Partners» и ее подразделения «Центр инклюзивного найма». В процессе участвуют HR-менеджер по инклюзивному найму, специалист по адаптации рабочих мест, модератор платформы и руководитель отдела подбора. HR-менеджер создает вакансии, специалист уточняет условия доступности, модератор проверяет корректность описаний, а руководитель оценивает отклики и результативность подбора.",
        "В терминологии проекта соискатель (candidate, applicant) — пользователь, который ищет работу и формирует заявку; вакансия (job vacancy, job posting) — описание рабочего места; работодатель (employer, company) — организация, публикующая вакансию; заявка (application, request) — набор выбранных вакансий и данных соискателя; профиль (candidate profile) — контактные и анкетные сведения пользователя. Эти понятия согласуются с терминологией исследований по inclusive employment и job matching [4].",
        "В статье Kuznetsov et al. 2025 описана production-ready система инклюзивного подбора работы, где центральными сущностями являются candidate, company, job position, readiness и compatibility score. Цель исследования — показать, что ML-система подбора может быть встроена в реальный процесс службы занятости. Среди заявленных результатов: F1-score 90,1 %, время ответа менее 100 мс и обработка 500 000 комбинаций candidate-company менее чем за 10 минут [4]. Для JobAbility из этой работы важен вывод: система должна хранить не только вакансию, но и признаки доступности рабочего места.",
        "В работе Kamruzzaman и Kim 2025 рассматривается disability disclosure — раскрытие сведений об инвалидности кандидата — и влияние этого фактора на fairness in candidate selection. Авторы показывают риск смещения в LLM-driven hiring: при прочих равных условиях кандидаты, явно указавшие отсутствие инвалидности, могут получать преимущество [5]. Поэтому JobAbility не выполняет автоматический отказ по данным о здоровье: профиль используется для оформления заявки, а финальное решение в процессе оставлено человеку-модератору.",
        "Для ручной аннотации терминов рекомендуется распечатать preprint arXiv:2508.11713. В нем удобно подписать перевод ключевых терминов: user — пользователь, candidate — соискатель, company — работодатель, job position — вакансия, application — заявка, accessibility/accommodation — доступность и адаптация рабочего места, matching — подбор.",
    ]
    for text in domain:
        p(doc, text)
    table(doc, "Таблица 1 — Термины предметной области с переводом и источниками", ["Термин RU", "English term", "Назначение в JobAbility", "Источник"], [
        ["Соискатель", "candidate / applicant", "Пользователь, выбирающий вакансии и формирующий заявку", "[4], [5]"],
        ["Работодатель", "employer / company", "Публикует вакансии и просматривает отклики", "[4]"],
        ["Вакансия", "job vacancy / job posting", "Рабочее место с зарплатой, графиком и условиями доступности", "[4]"],
        ["Заявка", "application / request", "Черновик или сформированное обращение соискателя", "[4]"],
        ["Раскрытие инвалидности", "disability disclosure", "Чувствительный признак, который нельзя использовать для дискриминации", "[5]"],
        ["Справедливость отбора", "fairness in candidate selection", "Требование не ухудшать положение кандидата из-за признака инвалидности", "[5]"],
    ], widths=[3.1, 3.6, 7.1, 1.8], font_size=9)

    p(doc, "Ключевой расчет предметной области — расчет суммарной ожидаемой зарплаты по заявке. Если заявка содержит n выбранных вакансий, а в каждой строке указано количество позиций qᵢ и зарплата вакансии sᵢ, итог определяется формулой (1).")
    formula(doc, "S = ∑ᵢ₌₁ⁿ(qᵢ × sᵢ)", "(1)")
    p(doc, "Для интерфейса также рассчитывается ориентировочная дата ответа. Она не является юридическим сроком, но помогает пользователю понимать ожидаемую длительность обработки заявки.")
    formula(doc, "D_response = D_formed + min(30, max(1, n) × 3) days", "(2)")
    p(doc, "Для оценки полноты описания вакансии используется показатель заполненности условий доступности рабочего места: отношение заполненных обязательных атрибутов к общему числу требуемых атрибутов.")
    formula(doc, "C_access = (m_filled / m_required) × 100%", "(3)")

    heading(doc, "Патентный поиск и конкурентные решения", 2)
    table(doc, "Таблица 2 — Результаты патентного поиска", ["№", "Источник", "Суть решения", "Релевантность", "Отличие JobAbility"], [
        ["1", "Google Patents: US20140279637A1", "Публикация по matching workflow между кандидатами и вакансиями", "Подтверждает типовую задачу сопоставления профиля и вакансии", "В JobAbility акцент сделан на доступности рабочих мест, заявке и модерации"],
        ["2", "Google Patents: search query «employment matching job seeker employer database»", "Класс решений job matching с профилями пользователей и работодателей", "Показывает зрелость рынка автоматизированного рекрутинга", "Система не заменяет решение человека, а поддерживает процесс и хранит признаки доступности"],
        ["3", "Google Patents: search query «disability employment accommodation job matching»", "Решения по подбору с учетом специальных условий и адаптации", "Близко к предметной области инклюзивного найма", "Реализован учебный web-сервис с открытым API, CI/CD и Kubernetes"],
    ], widths=[0.7, 2.9, 3.7, 3.6, 4.2], font_size=8)
    table(doc, "Таблица 3 — Конкурентные решения", ["Решение", "Сильные стороны", "Ограничение относительно JobAbility", "Что учтено в проекте"], [
        ["hh.ru", "Большая база вакансий, фильтры, личные кабинеты", "Нет фокуса на заявке для людей с ОВЗ и модерации условий доступности", "Каталог, карточка вакансии, кабинет работодателя"],
        ["Работа России", "Государственный портал занятости, связь с мерами поддержки", "Сложный универсальный сценарий, меньше учебной прозрачности API", "Ролевая модель и социальная направленность"],
        ["LinkedIn / Indeed", "Международный поиск, профили, рекомендации", "Инклюзивные признаки часто не являются центральной сущностью", "Публичный каталог и профиль пользователя"],
        ["AbilityJobs / GettingHired", "Специализация на employment for people with disabilities", "Зарубежный рынок и другая нормативная база", "Фокус на доступности, безопасном отклике и человеческой модерации"],
    ], widths=[2.7, 4.0, 4.6, 3.9], font_size=8)
    p(doc, "Вывод по патентному и конкурентному анализу: JobAbility не претендует на уникальность идеи поиска работы как таковой. Отличие проекта состоит в учебной, но законченной реализации полного цикла: от публикации доступной вакансии до формирования заявки, модерации, хранения медиа, мониторинга и автоматической доставки в Kubernetes.")

    if use_case.exists():
        picture(doc, use_case, "Рисунок 1 — Диаграмма прецедентов JobAbility", 15.5)
    if state.exists():
        picture(doc, state, "Рисунок 2 — Диаграмма состояний заявки", 15.5)
    add_landscape(doc)
    picture(doc, bpmn, "Рисунок 3 — Детализированная BPMN-диаграмма бизнес-процесса", 24.5)
    add_portrait(doc)

    heading(doc, "Архитектура")
    arch = [
        "Архитектура JobAbility построена как контейнеризованная web-система. Пользователь работает с React SPA, собранным Vite и размещенным в контейнере Nginx. Статический web-сервер отдает HTML, CSS и JavaScript, а запросы к /api, /admin, /swagger и /metrics передаются в backend-сервис.",
        "Backend реализован на Django 5.2 и Django REST Framework. Он отвечает за домены Auth, Vacancy, Application, ApplicationLine, Profile, EmployerResponses и System. В backend сосредоточены правила ролей, проверка входных данных, расчет заявки, создание черновика, модерация вакансий и заявок, а также выдача OpenAPI/Swagger.",
        "В Kubernetes приложение размещено в namespace jobability. Внутри кластера работают pods frontend, backend, postgres, redis, minio, prometheus, grafana и adminer. Traefik Ingress принимает HTTP-трафик на 80 порту и направляет его во frontend service. Backend обращается к PostgreSQL по TCP 5432, Redis по TCP 6379, MinIO по S3-compatible HTTP API на 9000 порту.",
        "CI/CD реализован в GitLab. Pipeline содержит стадии build, test, upload и deploy. На build собираются Docker-образы backend и frontend, на test запускаются Django-тесты в backend-образе, на upload образы импортируются в containerd-хранилище k3s, на deploy применяются Kubernetes-манифесты и обновляются deployment/backend и deployment/frontend.",
    ]
    for text in arch:
        p(doc, text)
    add_landscape(doc)
    picture(doc, deploy_img, "Рисунок 4 — Диаграмма развертывания JobAbility", 24.8)
    add_portrait(doc)
    picture(doc, er, "Рисунок 5 — ER-диаграмма основных таблиц", 15.5)
    table(doc, "Таблица 4 — Назначение таблиц", ["Таблица", "Назначение"], [
        ["auth_user", "Учетные записи Django: логин, email, пароль и базовая идентификация пользователя"],
        ["core_applicantprofile", "Анкетные данные соискателя: ФИО, телефон, город, возраст, пол, категория инвалидности"],
        ["core_vacancy", "Вакансии работодателей: название, компания, город, зарплата, описание, график, поддержка доступности, изображение"],
        ["core_application", "Заявки соискателей со статусным циклом DRAFT → FORMED → FINISHED/REJECTED"],
        ["core_applicationvacancy", "Строки заявки, связывающие заявку с выбранными вакансиями и количеством позиций"],
    ], widths=[4.2, 11.2], font_size=9)
    table(doc, "Таблица 5 — Описание колонок основных таблиц", ["Таблица", "Колонка", "Тип", "Описание"], [
        ["auth_user", "id", "int", "Первичный ключ пользователя"],
        ["auth_user", "username, email", "varchar", "Логин и электронная почта"],
        ["core_applicantprofile", "full_name, phone, city", "varchar", "Контактные данные соискателя"],
        ["core_applicantprofile", "age, gender, disability_category", "int/varchar", "Анкетные сведения для оформления заявки"],
        ["core_vacancy", "title, company, city", "varchar", "Основные сведения о вакансии"],
        ["core_vacancy", "salary, schedule, disability_support", "int/varchar", "Зарплата, график и условия доступности"],
        ["core_vacancy", "image", "varchar(100)", "Имя/путь медиафайла, обслуживаемого через MinIO"],
        ["core_application", "status", "varchar(16)", "DRAFT, FORMED, FINISHED, REJECTED, DELETED"],
        ["core_application", "formed_at, completed_at", "timestamptz", "Даты формирования и завершения заявки"],
        ["core_application", "total_salary", "int", "Итоговый расчет по строкам заявки"],
        ["core_applicationvacancy", "application_id, vacancy_id", "bigint FK", "Связь заявки и вакансии"],
        ["core_applicationvacancy", "qty, comment, is_main", "int/text/bool", "Количество, комментарий и признак основной вакансии"],
    ], widths=[3.2, 4.0, 3.0, 6.2], font_size=8)
    if class_img.exists():
        picture(doc, class_img, "Рисунок 6 — Диаграмма классов frontend/backend доменов", 15.5)
    p(doc, "Расчет аппаратных требований выполнен для 1000 зарегистрированных пользователей, пиковой активности 5 % и средней интенсивности 0,2 запроса в секунду на активного пользователя.")
    formula(doc, "RPS_peak = U × k_active × r_user = 1000 × 0.05 × 0.2 = 10 requests/s", "(4)")
    p(doc, "При расчетной стоимости backend-запроса 25 мс CPU и целевой утилизации 65 % минимальное количество ядер определяется формулой (5).")
    formula(doc, "CPU_cores = ceil(RPS_peak × t_cpu / ρ) = ceil(10 × 0.025 / 0.65) = 1", "(5)")
    p(doc, "С учетом PostgreSQL, Redis, MinIO, Prometheus, Grafana, k3s, containerd и резерва рекомендуется не менее 2 vCPU и 6-8 GiB RAM. Использованная виртуальная машина Ubuntu 22.04 с 7,1 GiB RAM и диском 23 GiB достаточна для демонстрационного контура.")
    formula(doc, "DB_year = U·B_user + V·B_vacancy + A·B_application + L·B_line", "(6)")
    p(doc, "При 1000 пользователях, 500 вакансиях, 1500 заявках и среднем числе 3 строк на заявку расчетный объем PostgreSQL до индексов составляет около 13 MiB, а с коэффициентом 3 на индексы и служебные данные — 40-60 MiB. Основной рост ожидается в MinIO: 500 изображений по 300 KiB занимают около 150 MiB, поэтому объектное хранилище выделено отдельным PVC.")

    heading(doc, "Алгоритмы")
    alg = [
        "Основной алгоритм представлен последовательностью HTTP-запросов между браузером, frontend-страницами, API-клиентами и backend-доменами. Сначала пользователь выполняет POST /api/users/login/ или POST /api/users/register/. Backend проверяет данные, создает или открывает сессию и возвращает сведения о роли пользователя. После обновления страницы frontend вызывает GET /api/users/me/, чтобы восстановить состояние текущего пользователя.",
        "Гость получает список вакансий через GET /api/vacancies/. В запрос могут передаваться search, min_price, max_price и другие фильтры. При открытии карточки выполняется GET /api/vacancies/{id}/. Если пользователь авторизован как соискатель, интерфейс дополнительно вызывает GET /api/applications/cart/. Если черновика нет, backend возвращает application_id = null и items_count = 0, что соответствует пустой иконке заявки.",
        "Добавление вакансии в заявку выполняется POST /api/application-lines/ с vacancy_id и qty. Backend находит или создает черновую заявку DRAFT, добавляет строку и возвращает обновленные данные. Редактирование строки выполняется PUT /api/application-lines/ с vacancy_id, qty, comment, is_main и order_index; удаление строки — DELETE /api/application-lines/. Просмотр черновой заявки выполняется GET /api/applications/{id}/.",
        "Формирование заявки выполняется PUT /api/applications/{id}/form/. Backend проверяет профиль, наличие строк, рассчитывает total_salary, фиксирует formed_at и переводит заявку в статус FORMED. Модератор получает список заявок через GET /api/applications/ с фильтрами, открывает заявку GET /api/applications/{id}/ и завершает ее PUT /api/applications/{id}/moderate/ с action = finish или reject. Работодатель получает список откликов через GET /api/applications/employer-responses/.",
    ]
    for text in alg:
        p(doc, text)
    if seq_part_1.exists() and seq_part_2.exists():
        picture(doc, seq_part_1, "Рисунок 7 — Диаграмма последовательности HTTP-запросов (часть 1)", 15.0)
        picture(doc, seq_part_2, "Рисунок 8 — Диаграмма последовательности HTTP-запросов (часть 2)", 15.0)
    elif seq_img.exists():
        picture(doc, seq_img, "Рисунок 7 — Диаграмма последовательности HTTP-запросов", 15.0)

    heading(doc, "Описание интерфейса")
    iface = [
        "Интерфейс JobAbility реализован как адаптивное SPA-приложение. Навигационная панель содержит переходы на главную страницу, каталог вакансий, Swagger, административные разделы и действия входа/регистрации. На мобильных экранах блоки перестраиваются вертикально, карточки вакансий занимают всю ширину, а формы сохраняют крупные поля ввода.",
        "Главная страница объясняет назначение сервиса: система трудоустройства для людей с ограниченными возможностями. Страница регистрации позволяет выбрать роль соискателя или работодателя, указать имя, фамилию, username, email и пароль. Страница входа выполняет авторизацию и переводит пользователя в сценарий, соответствующий его роли.",
        "Каталог вакансий отображает карточки с названием, компанией, городом, зарплатой, графиком и описанием поддержки доступности. Карточка вакансии показывает подробное описание и позволяет добавить вакансию в заявку. Кабинет соискателя содержит профиль и переход к списку заявок. Страница заявки позволяет редактировать строки, комментарии, порядок и формировать заявку.",
        "Кабинет работодателя используется для создания вакансий и просмотра статуса модерации. Страница откликов работодателя показывает заявки, в которых присутствуют его вакансии. Кабинет модератора содержит список вакансий на проверке и список заявок; модератор может одобрять или отклонять вакансии, завершать или отклонять заявки.",
    ]
    for text in iface:
        p(doc, text)
    table(doc, "Таблица 6 — Окна интерфейса и действия пользователей", ["Окно", "Маршрут", "Действия", "API"], [
        ["Главная", "/", "Переход к каталогу, входу, регистрации", "GET /"],
        ["Регистрация", "/register", "Создание аккаунта соискателя или работодателя", "POST /api/users/register/"],
        ["Вход", "/login", "Авторизация и восстановление сессии", "POST /api/users/login/, GET /api/users/me/"],
        ["Вакансии", "/vacancies", "Поиск, фильтрация, открытие карточек", "GET /api/vacancies/"],
        ["Карточка вакансии", "/vacancies/:id", "Просмотр и добавление в заявку", "GET /api/vacancies/{id}/, POST /api/application-lines/"],
        ["Заявки", "/applications", "Список, фильтры, расчет сумм", "GET /api/applications/"],
        ["Детальная заявка", "/applications/:id", "Редактирование, формирование, модерация", "GET/PUT /api/applications/{id}/"],
        ["Кабинет соискателя", "/cabinet/applicant", "Профиль, пароль, переход к заявкам", "GET/PUT /api/users/profile/"],
        ["Кабинет работодателя", "/cabinet/employer", "Создание вакансий и просмотр статусов", "POST /api/vacancies/, GET /api/vacancies/mine/"],
        ["Отклики работодателя", "/cabinet/employer/responses", "Просмотр заявок по вакансиям работодателя", "GET /api/applications/employer-responses/"],
        ["Кабинет модератора", "/cabinet/moderator", "Модерация вакансий и заявок", "GET /api/vacancies/pending/, PUT /moderate/"],
        ["Swagger", "/swagger/", "Просмотр API-документации", "GET /swagger/"],
    ], widths=[3.1, 3.5, 5.3, 5.4], font_size=8)

    heading(doc, "Заключение")
    for text in [
        "В результате работы разработана и развернута система JobAbility — web-сервис трудоустройства для людей с ограниченными возможностями. Реализованы роли гостя, соискателя, работодателя и модератора; каталог вакансий; карточка вакансии; профиль соискателя; черновик и формирование заявки; модерация вакансий и заявок; кабинеты пользователей.",
        "Подготовлена инфраструктура: Docker-образы backend и frontend, Kubernetes-манифесты для frontend, backend, PostgreSQL, Redis, MinIO, Prometheus, Grafana и Adminer, Ingress через Traefik, PVC для постоянных данных, Secret и ConfigMap. Настроен GitLab Runner с shell executor и pipeline build-test-upload-deploy. Тестовая стадия запускает Django-тесты перед доставкой.",
        "Итоговый результат подтвержден успешным pipeline и работающим приложением на виртуальной машине. Проект опубликован в GitLab: https://gitlab.com/xflame0xx1/jobability. Демо через временный Cloudflare Tunnel использовалось для проверки публичного доступа; локальный адрес стенда: http://192.168.56.19/.",
    ]:
        p(doc, text)

    heading(doc, "Список использованных источников")
    sources = [
        "World Health Organization. Disability: fact sheet. URL: https://www.who.int/news-room/fact-sheets/detail/disability-and-health (дата обращения: 29.05.2026).",
        "International Labour Organization. Disability inclusion at work. URL: https://www.ilo.org/topics/disability-and-work (дата обращения: 29.05.2026).",
        "Федеральный закон от 24.11.1995 № 181-ФЗ «О социальной защите инвалидов в Российской Федерации». URL: https://www.consultant.ru/document/cons_doc_LAW_8559/ (дата обращения: 29.05.2026).",
        "Kuznetsov O., Melchiori M., Frontoni E., Arnesano M. A Production-Ready Machine Learning System for Inclusive Employment: Requirements Engineering and Implementation of AI-Driven Disability Job Matching Platform. arXiv:2508.11713, 2025. URL: https://arxiv.org/abs/2508.11713.",
        "Kamruzzaman M., Kim G. L. The Impact of Disability Disclosure on Fairness and Bias in LLM-Driven Candidate Selection. arXiv:2506.00256, 2025. URL: https://arxiv.org/abs/2506.00256.",
        "Soiffer N. et al. Barriers to Employment: The Deaf Multimedia Authoring Tax. arXiv:2505.01030, 2025. URL: https://arxiv.org/abs/2505.01030.",
        "Django Software Foundation. Django documentation 5.2. URL: https://docs.djangoproject.com/en/5.2/.",
        "Django REST framework. API Guide. URL: https://www.django-rest-framework.org/api-guide/.",
        "React documentation. URL: https://react.dev/.",
        "Kubernetes Documentation. Concepts: workloads, services, ingress. URL: https://kubernetes.io/docs/concepts/.",
        "Docker Documentation. Build and Dockerfile reference. URL: https://docs.docker.com/build/.",
        "PostgreSQL Documentation. URL: https://www.postgresql.org/docs/.",
        "GitLab Documentation. CI/CD pipelines. URL: https://docs.gitlab.com/ci/.",
        "MinIO Documentation. URL: https://min.io/docs/minio/kubernetes/upstream/.",
        "Redis Documentation. URL: https://redis.io/docs/latest/.",
        "Google Patents. Patent publication US20140279637A1 and patent search results for employment matching systems. URL: https://patents.google.com/.",
        "HeadHunter. URL: https://hh.ru/.",
        "Портал «Работа России». URL: https://trudvsem.ru/.",
        "AbilityJobs. URL: https://www.abilityjobs.com/.",
    ]
    for i, src in enumerate(sources, 1):
        p(doc, f"{i}. {src}")

    heading(doc, "Приложение А. Техническое задание")
    tz = [
        "1. Цель разработки — создать итоговую систему JobAbility для поддержки трудоустройства людей с ограниченными возможностями, включая публикацию доступных вакансий, формирование заявок, модерацию и демонстрационное развертывание в Kubernetes.",
        "2. Назначение системы — предоставить гостям публичный каталог вакансий, соискателям инструмент подбора и оформления заявки, работодателям кабинет публикации вакансий, модератору инструмент контроля качества данных и завершения заявок.",
        "3. Задачи разработки: 3.1 реализовать backend-домены Auth, Vacancy, Application, Profile на Django 5.2 и DRF; 3.2 реализовать REST API с OpenAPI/Swagger; 3.3 реализовать роли и session-based authentication; 3.4 реализовать React/Vite frontend с адаптивными страницами; 3.5 реализовать черновик заявки и расчет total_salary; 3.6 реализовать кабинеты работодателя и модератора; 3.7 добавить Django-тесты и CI test stage; 3.8 контейнеризовать backend/frontend и развернуть в k3s; 3.9 подготовить нативный сценарий/desktop-демо на базе web-интерфейса при необходимости; 3.10 подготовить демо GitHub Pages для статического описания проекта; 3.11 оформить документацию, диаграммы, Swagger и РПЗ; 3.12 вести репозиторий Git/GitLab с pipeline build-test-upload-deploy.",
        "4. Функциональные требования. 4.1 HTTP-методы системы должны включать GET, POST, PUT и DELETE для доменов Auth, Vacancy, Profile, Application, ApplicationLine, EmployerResponses и System. 4.2 Меню должно содержать переходы на главную страницу, вакансии, Swagger, вход, регистрацию, кабинеты и заявки в зависимости от роли. 4.3 Гость должен иметь страницу регистрации с вызовом POST /api/users/register/. 4.4 Гость и пользователь должны иметь страницу аутентификации с POST /api/users/login/ и GET /api/users/me/. 4.5 Гость должен видеть список услуг/вакансий через GET /api/vacancies/. 4.6 Гость или соискатель должен видеть одну услугу/вакансию через GET /api/vacancies/{id}/, а соискатель должен добавлять ее в заявку через POST /api/application-lines/. 4.7 Создатель заявки должен иметь страницу заявки и список заявок с GET /api/applications/, GET/PUT /api/applications/{id}/, PUT /api/applications/{id}/form/. 4.8 Модератор должен иметь дополнительные функции списка заявок и вакансий: GET /api/vacancies/pending/, PUT /api/vacancies/{id}/moderate/, PUT /api/applications/{id}/moderate/.",
        "5. Требования к аппаратному обеспечению сервера: минимум 2 vCPU, 6-8 GiB RAM, 23 GiB диска для демонстрационного стенда; для роста медиафайлов предусмотреть отдельный PVC MinIO. Клиент: любой ПК или смартфон с современным браузером, экраном от 360 px по ширине и доступом к HTTP/HTTPS.",
        "6. Требования к программному обеспечению сервера: Ubuntu 22.04 LTS, Docker Engine 29.2.1, k3s v1.35.5+k3s1, containerd 2.2.3-k3s1, GitLab Runner 19.0.0, Python 3.13, Django 5.2, Django REST Framework, PostgreSQL, Redis, MinIO, Nginx, Prometheus, Grafana, Adminer. Клиент: Chrome/Edge/Firefox актуальной версии, поддержка JavaScript ES2022 и CSS Grid/Flexbox.",
    ]
    for text in tz:
        p(doc, text)

    add_landscape(doc)
    heading(doc, "Приложение Б. Методы веб-сервиса")
    api_rows = [
        ["1", "Auth", "GET", "/api/users/me/", "Текущий пользователь", "-", "id:int, username:str, role:str, email:str"],
        ["2", "Auth", "POST", "/api/users/register/", "Регистрация", "username:str, password:str, first_name:str, last_name:str, email:str, role:str", "User"],
        ["3", "Auth", "POST", "/api/users/login/", "Вход", "username:str, password:str", "User + session"],
        ["4", "Auth", "POST", "/api/users/logout/", "Выход", "-", "message:str"],
        ["5", "Auth", "PUT", "/api/users/password/", "Смена пароля", "old_password:str, new_password:str", "message:str"],
        ["6", "Vacancy", "GET", "/api/vacancies/", "Список вакансий", "search:str, city:str, salary:int", "Vacancy[]"],
        ["7", "Vacancy", "POST", "/api/vacancies/", "Создание вакансии", "title, company, city, salary, description, image:file", "Vacancy"],
        ["8", "Vacancy", "GET", "/api/vacancies/{id}/", "Карточка вакансии", "id:int", "Vacancy"],
        ["9", "Vacancy", "GET", "/api/vacancies/mine/", "Мои вакансии", "-", "Vacancy[]"],
        ["10", "Vacancy", "GET", "/api/vacancies/pending/", "Вакансии на модерации", "-", "Vacancy[]"],
        ["11", "Vacancy", "PUT", "/api/vacancies/{id}/moderate/", "Модерация вакансии", "action:str, note:str", "Vacancy"],
        ["12", "Profile", "GET", "/api/users/profile/", "Профиль", "-", "ApplicantProfile"],
        ["13", "Profile", "PUT", "/api/users/profile/", "Обновить профиль", "full_name, phone, city, age, gender, disability_category", "ApplicantProfile"],
        ["14", "Application", "GET", "/api/applications/cart/", "Пустая/текущая корзина", "-", "application_id:int|null, items_count:int"],
        ["15", "ApplicationLine", "POST", "/api/application-lines/", "Добавить вакансию", "vacancy_id:int, qty:int", "ApplicationLine"],
        ["16", "ApplicationLine", "PUT", "/api/application-lines/", "Изменить строку", "vacancy_id:int, qty:int, comment:str, is_main:bool", "ApplicationLine"],
        ["17", "ApplicationLine", "DELETE", "/api/application-lines/", "Удалить строку", "vacancy_id:int", "204 No Content"],
        ["18", "Application", "GET", "/api/applications/", "Список заявок", "status:str, date_from:date", "Application[]"],
        ["19", "Application", "GET", "/api/applications/{id}/", "Детали заявки", "id:int", "ApplicationDetail"],
        ["20", "Application", "PUT", "/api/applications/{id}/", "Редактировать заявку", "profile/contact fields", "ApplicationDetail"],
        ["21", "Application", "PUT", "/api/applications/{id}/form/", "Сформировать заявку", "id:int", "ApplicationDetail"],
        ["22", "Application", "PUT", "/api/applications/{id}/moderate/", "Завершить/отклонить", "action:finish|reject, moderator_note:str", "ApplicationDetail"],
        ["23", "Application", "DELETE", "/api/applications/{id}/delete/", "Удалить черновик", "id:int", "204 No Content"],
        ["24", "Employer", "GET", "/api/applications/employer-responses/", "Отклики работодателя", "-", "Application[]"],
        ["25", "System", "GET", "/api/schema/", "OpenAPI schema", "-", "OpenAPI JSON"],
        ["26", "System", "GET", "/swagger/", "Swagger UI", "-", "HTML"],
        ["27", "System", "GET", "/metrics/", "Метрики Prometheus", "-", "text/plain"],
    ]
    grouped: dict[str, list[list[str]]] = {}
    for row in api_rows:
        grouped.setdefault(row[1], []).append([row[0], row[2], row[3], row[4], row[5], row[6]])
    n = 1
    for domain, rows in grouped.items():
        p(doc, f"Домен {domain}.", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        table(doc, f"Таблица Б.{n} — Методы домена {domain}", ["№", "Метод", "URL", "Описание", "Входные данные", "Выходные данные"], rows, widths=[0.8, 1.2, 4.0, 4.2, 6.0, 5.2], font_size=8)
        n += 1

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
