from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "rpz"
IMG_DIR = OUT_DIR / "images"
DOCX_PATH = OUT_DIR / "RPZ_JobAbility.docx"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


F_TITLE = font(34, True)
F_H = font(24, True)
F_B = font(17, True)
F = font(15)
F_S = font(12)


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        candidate = f"{line} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=fnt)[2] <= width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: Iterable[str] = (),
    fill: str = "#FFFFFF",
    outline: str = "#3056D3",
    title_color: str = "#0F172A",
    radius: int = 18,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    draw.text((x1 + 18, y1 + 14), title, font=F_B, fill=title_color)
    y = y1 + 45
    for line in body:
        for part in wrap(draw, line, F, x2 - x1 - 36):
            draw.text((x1 + 18, y), part, font=F, fill="#334155")
            y += 22


def arrow(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    color: str = "#2563EB",
    width: int = 3,
    text: str | None = None,
    text_pos: tuple[int, int] | None = None,
) -> None:
    draw.line(points, fill=color, width=width, joint="curve")
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    dx, dy = x2 - x1, y2 - y1
    if abs(dx) > abs(dy):
        if dx >= 0:
            head = [(x2, y2), (x2 - 12, y2 - 7), (x2 - 12, y2 + 7)]
        else:
            head = [(x2, y2), (x2 + 12, y2 - 7), (x2 + 12, y2 + 7)]
    else:
        if dy >= 0:
            head = [(x2, y2), (x2 - 7, y2 - 12), (x2 + 7, y2 - 12)]
        else:
            head = [(x2, y2), (x2 - 7, y2 + 12), (x2 + 7, y2 + 12)]
    draw.polygon(head, fill=color)
    if text and text_pos:
        tw = draw.textbbox((0, 0), text, font=F_S)[2] + 14
        draw.rounded_rectangle(
            (text_pos[0] - 6, text_pos[1] - 4, text_pos[0] + tw, text_pos[1] + 17),
            radius=8,
            fill="#F8FAFC",
            outline="#CBD5E1",
        )
        draw.text(text_pos, text, font=F_S, fill="#475569")


def save_canvas(path: Path, size: tuple[int, int], title: str) -> ImageDraw.ImageDraw:
    img = Image.new("RGB", size, "#F8FAFC")
    draw = ImageDraw.Draw(img)
    draw.text((size[0] // 2, 26), title, font=F_TITLE, anchor="ma", fill="#0F172A")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return draw


def make_use_case(path: Path) -> None:
    img = Image.new("RGB", (1900, 1200), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    draw.text((950, 35), "JobAbility - диаграмма прецедентов", font=F_TITLE, anchor="ma", fill="#0F172A")
    draw.rounded_rectangle((360, 120, 1540, 1080), 28, fill="#FFFFFF", outline="#CBD5E1", width=3)
    draw.text((950, 150), "Система трудоустройства для людей с ограниченными возможностями", font=F_H, anchor="ma", fill="#1E3A8A")

    actors = {
        "Гость": (140, 260),
        "Соискатель": (135, 610),
        "Работодатель": (1710, 380),
        "Модератор": (1710, 760),
    }
    for name, (x, y) in actors.items():
        draw.ellipse((x - 28, y - 60, x + 28, y - 4), outline="#0F172A", width=3)
        draw.line((x, y - 4, x, y + 78), fill="#0F172A", width=3)
        draw.line((x - 48, y + 20, x + 48, y + 20), fill="#0F172A", width=3)
        draw.line((x, y + 78, x - 42, y + 138), fill="#0F172A", width=3)
        draw.line((x, y + 78, x + 42, y + 138), fill="#0F172A", width=3)
        draw.text((x, y + 155), name, font=F_B, anchor="ma", fill="#0F172A")

    use_cases = [
        ("Регистрация и вход", (520, 250)),
        ("Просмотр и фильтрация вакансий", (520, 420)),
        ("Просмотр карточки вакансии", (520, 590)),
        ("Добавление вакансии в заявку", (900, 360)),
        ("Заполнение профиля и черновика", (900, 540)),
        ("Формирование заявки", (900, 720)),
        ("Просмотр статусов заявок", (900, 900)),
        ("Создание вакансии", (1280, 320)),
        ("Просмотр откликов", (1280, 500)),
        ("Модерация вакансий", (1280, 710)),
        ("Завершение или отклонение заявки", (1280, 890)),
    ]
    centers = {}
    for text, (x, y) in use_cases:
        centers[text] = (x, y)
        draw.ellipse((x - 175, y - 48, x + 175, y + 48), fill="#EEF2FF", outline="#4F46E5", width=3)
        for i, line in enumerate(wrap(draw, text, F_B, 280)):
            draw.text((x, y - 17 + i * 22), line, font=F_B, anchor="ma", fill="#1E1B4B")

    links = [
        ("Гость", "Регистрация и вход"), ("Гость", "Просмотр и фильтрация вакансий"), ("Гость", "Просмотр карточки вакансии"),
        ("Соискатель", "Добавление вакансии в заявку"), ("Соискатель", "Заполнение профиля и черновика"),
        ("Соискатель", "Формирование заявки"), ("Соискатель", "Просмотр статусов заявок"),
        ("Работодатель", "Создание вакансии"), ("Работодатель", "Просмотр откликов"),
        ("Модератор", "Модерация вакансий"), ("Модератор", "Завершение или отклонение заявки"),
    ]
    for actor, uc in links:
        ax, ay = actors[actor]
        ux, uy = centers[uc]
        start = (ax + 70 if ax < 900 else ax - 70, ay + 55)
        end = (ux - 175 if ax < 900 else ux + 175, uy)
        draw.line((start, end), fill="#64748B", width=2)
    img.save(path)


def make_state(path: Path) -> None:
    img = Image.new("RGB", (1800, 1050), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    draw.text((900, 35), "JobAbility - диаграмма состояний заявки", font=F_TITLE, anchor="ma", fill="#0F172A")
    states = {
        "Начало": (150, 475, 250, 575),
        "DRAFT\nЧерновик": (420, 430, 720, 620),
        "FORMED\nСформирована": (900, 430, 1220, 620),
        "FINISHED\nЗавершена": (1440, 230, 1720, 400),
        "REJECTED\nОтклонена": (1440, 650, 1720, 820),
        "DELETED\nУдалена": (560, 770, 820, 930),
    }
    for name, xy in states.items():
        fill = "#FFFFFF"
        outline = "#2563EB"
        if name == "Начало":
            draw.ellipse(xy, fill="#111827", outline="#111827")
            continue
        if name.startswith("FINISHED"):
            fill, outline = "#DCFCE7", "#16A34A"
        elif name.startswith("REJECTED"):
            fill, outline = "#FEE2E2", "#DC2626"
        elif name.startswith("DELETED"):
            fill, outline = "#F1F5F9", "#64748B"
        draw.rounded_rectangle(xy, 20, fill=fill, outline=outline, width=3)
        cx = (xy[0] + xy[2]) // 2
        cy = (xy[1] + xy[3]) // 2
        for i, line in enumerate(name.split("\n")):
            draw.text((cx, cy - 18 + i * 28), line, font=F_B, anchor="ma", fill="#0F172A")
    arrow(draw, [(250, 525), (420, 525)], text="создание черновика", text_pos=(285, 495))
    arrow(draw, [(720, 525), (900, 525)], text="form()", text_pos=(770, 495))
    arrow(draw, [(1220, 505), (1340, 505), (1340, 315), (1440, 315)], color="#16A34A", text="moderate(finish)", text_pos=(1260, 285))
    arrow(draw, [(1220, 555), (1340, 555), (1340, 735), (1440, 735)], color="#DC2626", text="moderate(reject)", text_pos=(1260, 705))
    arrow(draw, [(590, 620), (590, 770)], color="#64748B", text="delete()", text_pos=(605, 680))
    box(draw, (180, 790, 470, 930), "Ограничения", [
        "Из черновика можно сформировать или удалить заявку.",
        "После завершения/отклонения изменение статуса закрыто.",
    ], fill="#FFFFFF", outline="#CBD5E1")
    img.save(path)


def make_er(path: Path) -> None:
    img = Image.new("RGB", (2200, 1450), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    draw.text((1100, 35), "JobAbility - ER-диаграмма данных", font=F_TITLE, anchor="ma", fill="#0F172A")

    entities = {
        "auth_user": ((120, 180, 520, 405), ["id PK", "username", "password", "email", "first_name, last_name"]),
        "core_useraccount": ((760, 180, 1160, 380), ["id PK", "user_id FK", "role", "created_at"]),
        "core_applicantprofile": ((760, 520, 1160, 805), ["id PK", "user_id FK", "full_name", "phone, city, age", "gender", "disability_category"]),
        "core_vacancy": ((1420, 180, 2060, 565), ["id PK", "title, company, city", "salary", "description", "creator_id FK", "moderator_id FK", "moderation_status", "image, video", "schedule, disability_support"]),
        "core_application": ((760, 980, 1320, 1325), ["id PK", "creator_id FK", "moderator_id FK", "applicant_id FK", "status", "created_at, formed_at, completed_at", "contact_email, cover_letter", "estimated_response_date", "total_salary"]),
        "core_applicationvacancy": ((1460, 900, 2060, 1245), ["id PK", "application_id FK", "vacancy_id FK", "qty", "comment", "is_main", "order_index", "line_salary_total"]),
    }
    for name, (xy, fields) in entities.items():
        box(draw, xy, name, fields, fill="#FFFFFF", outline="#0EA5E9")

    arrow(draw, [(520, 275), (760, 275)], color="#475569", text="1 : 1", text_pos=(610, 245))
    arrow(draw, [(520, 330), (640, 330), (640, 650), (760, 650)], color="#475569", text="1 : 0..1", text_pos=(565, 610))
    arrow(draw, [(520, 360), (600, 360), (600, 1090), (760, 1090)], color="#475569", text="1 : N creator", text_pos=(545, 1040))
    arrow(draw, [(1160, 660), (1225, 660), (1225, 1140), (760, 1140)], color="#475569", text="1 : N applicant", text_pos=(1240, 900))
    arrow(draw, [(1320, 1110), (1460, 1110)], color="#475569", text="1 : N", text_pos=(1360, 1080))
    arrow(draw, [(1760, 900), (1760, 565)], color="#475569", text="N : 1", text_pos=(1780, 720))
    arrow(draw, [(1160, 275), (1420, 275)], color="#475569", text="creator / moderator", text_pos=(1225, 245))
    img.save(path)


def make_bpmn(path: Path) -> None:
    img = Image.new("RGB", (2600, 1450), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    draw.text((1300, 35), "JobAbility - BPMN бизнес-процесс трудоустройства", font=F_TITLE, anchor="ma", fill="#0F172A")
    lanes = [
        ("Соискатель", 130, 440, "#EFF6FF"),
        ("Работодатель", 440, 750, "#F0FDF4"),
        ("Сервис JobAbility / модератор", 750, 1260, "#FFF7ED"),
    ]
    for title, y1, y2, fill in lanes:
        draw.rounded_rectangle((80, y1, 2520, y2), 18, fill=fill, outline="#CBD5E1", width=2)
        draw.text((110, (y1 + y2) // 2), title, font=F_B, anchor="lm", fill="#0F172A")
    steps = [
        ("start", "Старт", (300, 265), "circle"),
        ("browse", "Просмотреть вакансии", (520, 265), "task"),
        ("add", "Добавить вакансию\nв заявку", (820, 265), "task"),
        ("profile", "Заполнить профиль\nи черновик", (1120, 265), "task"),
        ("form", "Сформировать\nзаявку", (1420, 265), "task"),
        ("track", "Отслеживать\nрезультат", (2220, 265), "task"),
        ("createVacancy", "Создать вакансию", (520, 585), "task"),
        ("viewResponses", "Просмотреть\nотклики", (1890, 585), "task"),
        ("validate", "Проверить данные\nи права доступа", (820, 900), "task"),
        ("moderateVacancy", "Модерация\nвакансии", (1120, 900), "task"),
        ("calc", "Рассчитать сумму\nи дату ответа", (1420, 900), "task"),
        ("moderateApp", "Завершить или\nотклонить заявку", (1890, 900), "task"),
        ("end", "Финиш", (2220, 900), "circle"),
    ]
    centers = {}
    for _, label, (x, y), kind in steps:
        centers[label] = (x, y)
        if kind == "circle":
            draw.ellipse((x - 45, y - 45, x + 45, y + 45), fill="#FFFFFF", outline="#0F172A", width=4)
            draw.text((x, y), label, font=F_B, anchor="ma", fill="#0F172A")
        else:
            box(draw, (x - 120, y - 55, x + 120, y + 55), label.replace("\n", " "), [], fill="#FFFFFF", outline="#2563EB")
            # overwrite with centered multiline text
            draw.rounded_rectangle((x - 120, y - 55, x + 120, y + 55), 18, fill="#FFFFFF", outline="#2563EB", width=3)
            for i, line in enumerate(label.split("\n")):
                draw.text((x, y - 18 + i * 26), line, font=F_B, anchor="ma", fill="#0F172A")

    def a(fr: tuple[int, int], to: tuple[int, int], label: str = "", color: str = "#2563EB"):
        arrow(draw, [fr, to], color=color, text=label or None, text_pos=((fr[0] + to[0]) // 2 - 40, (fr[1] + to[1]) // 2 - 35) if label else None)

    a((345, 265), (400, 265))
    a((640, 265), (700, 265))
    a((940, 265), (1000, 265))
    a((1240, 265), (1300, 265))
    arrow(draw, [(1540, 265), (1680, 265), (1680, 900), (1300, 900)], text="HTTP PUT /form", text_pos=(1560, 555))
    arrow(draw, [(640, 585), (1120, 585), (1120, 845)], color="#16A34A", text="новая вакансия", text_pos=(780, 555))
    arrow(draw, [(1240, 900), (1300, 900)], color="#F97316", text="одобрена", text_pos=(1230, 865))
    arrow(draw, [(1540, 900), (1770, 900)], color="#F97316", text="заявка", text_pos=(1615, 865))
    arrow(draw, [(2010, 900), (2175, 900)], color="#F97316", text="решение", text_pos=(2050, 865))
    arrow(draw, [(1890, 845), (1890, 640)], color="#16A34A", text="видит заявки", text_pos=(1905, 730))
    arrow(draw, [(2010, 585), (2220, 585), (2220, 320)], color="#16A34A", text="статус", text_pos=(2100, 555))
    img.save(path)


def make_er(path: Path) -> None:
    img = Image.new("RGB", (1900, 1050), "#F7F8FA")
    draw = ImageDraw.Draw(img)
    draw.text((950, 30), "JobAbility - ER-диаграмма основных таблиц", font=F_TITLE, anchor="ma", fill="#0F172A")

    def entity(x: int, y: int, w: int, title: str, fields: list[tuple[str, str]]) -> tuple[int, int, int, int]:
        row_h = 34
        h = 46 + row_h * len(fields)
        draw.rounded_rectangle((x + 8, y + 8, x + w + 8, y + h + 8), 8, fill="#E5E7EB")
        draw.rounded_rectangle((x, y, x + w, y + h), 8, fill="#F3F4F6", outline="#D1D5DB", width=2)
        draw.rounded_rectangle((x, y, x + w, y + 42), 8, fill="#176496", outline="#176496")
        draw.rectangle((x, y + 28, x + w, y + 42), fill="#176496")
        draw.text((x + 14, y + 11), title, font=F_B, fill="#FFFFFF")
        yy = y + 54
        for name, typ in fields:
            draw.text((x + 16, yy), name + (" key" if name == "id" else ""), font=F, fill="#111827")
            draw.text((x + w - 16, yy), typ, font=F, fill="#4B5563", anchor="ra")
            yy += row_h
        return (x, y, x + w, y + h)

    entity(70, 695, 300, "auth_user", [
        ("id", "int"),
        ("username", "varchar"),
        ("email", "varchar"),
    ])
    entity(515, 95, 405, "core_applicantprofile", [
        ("id", "bigint"),
        ("full_name", "varchar(255)"),
        ("phone", "varchar(64)"),
        ("city", "varchar(255)"),
        ("age", "int"),
        ("gender", "varchar(16)"),
        ("disability_category", "varchar(16)"),
        ("user_id", "int NN"),
    ])
    entity(1110, 180, 430, "core_application", [
        ("id", "bigint"),
        ("status", "varchar(16)"),
        ("created_at", "timestamptz"),
        ("formed_at", "timestamptz"),
        ("completed_at", "timestamptz"),
        ("total_salary", "int"),
        ("applicant_id", "bigint"),
        ("creator_id", "int NN"),
    ])
    entity(1620, 180, 360, "core_applicationvacancy", [
        ("id", "bigint"),
        ("qty", "int"),
        ("comment", "text"),
        ("is_main", "boolean"),
        ("order_index", "int"),
        ("application_id", "bigint NN"),
        ("vacancy_id", "bigint NN"),
    ])
    entity(880, 635, 430, "core_vacancy", [
        ("id", "bigint"),
        ("title", "varchar(255)"),
        ("company", "varchar(255)"),
        ("city", "varchar(255)"),
        ("salary", "int"),
        ("description", "text"),
        ("is_active", "boolean"),
        ("image", "varchar(100)"),
        ("disability_support", "varchar(255)"),
        ("schedule", "varchar(255)"),
    ])

    line_color = "#B7BEC7"
    arrow(draw, [(370, 755), (470, 755), (470, 280), (515, 280)], color=line_color, width=3)
    arrow(draw, [(370, 795), (760, 795), (760, 650), (880, 650)], color=line_color, width=3)
    arrow(draw, [(920, 315), (1110, 315)], color=line_color, width=3)
    arrow(draw, [(1540, 315), (1620, 315)], color=line_color, width=3)
    arrow(draw, [(1620, 525), (1480, 525), (1480, 720), (1310, 720)], color=line_color, width=3)
    arrow(draw, [(370, 835), (520, 835), (520, 460), (1110, 460)], color=line_color, width=3)
    img.save(path)


def make_bpmn(path: Path) -> None:
    img = Image.new("RGB", (3600, 1850), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((1800, 35), "JobAbility - детализированная BPMN-диаграмма процесса", font=F_TITLE, anchor="ma", fill="#0F172A")
    lanes = [
        ("Соискатель", 130, 580, "#F8FAFC"),
        ("Работодатель", 580, 930, "#F8FAFC"),
        ("Модератор", 930, 1280, "#F8FAFC"),
        ("Backend JobAbility", 1280, 1720, "#F8FAFC"),
    ]
    for title, y1, y2, fill in lanes:
        draw.rectangle((90, y1, 3510, y2), fill=fill, outline="#111827", width=2)
        draw.rectangle((90, y1, 155, y2), fill="#FFFFFF", outline="#111827", width=2)
        draw.text((122, (y1 + y2) // 2), title, font=F_B, anchor="mm", fill="#0F172A")

    def task(x, y, text, w=250, h=86):
        draw.rounded_rectangle((x - w // 2, y - h // 2, x + w // 2, y + h // 2), 16, fill="#FFFFFF", outline="#111827", width=3)
        lines = wrap(draw, text, F, w - 28)
        start = y - (len(lines) - 1) * 11
        for i, line in enumerate(lines):
            draw.text((x, start + i * 22), line, font=F, anchor="ma", fill="#111827")

    def gateway(x, y, text):
        draw.polygon([(x, y - 42), (x + 42, y), (x, y + 42), (x - 42, y)], fill="#FFFFFF", outline="#111827")
        draw.text((x, y + 62), text, font=F_S, anchor="ma", fill="#111827")

    def event(x, y, text, end=False):
        draw.ellipse((x - 34, y - 34, x + 34, y + 34), fill="#FFFFFF", outline="#111827", width=4 if end else 3)
        draw.text((x, y + 50), text, font=F_S, anchor="ma", fill="#111827")

    nodes = {
        "start": (240, 350), "reg": (430, 350), "login": (700, 350), "list": (980, 350),
        "detail": (1260, 350), "cart": (1540, 350), "profile": (1820, 350),
        "draft": (2100, 350), "valid": (2380, 350), "formed": (2660, 350), "watch": (3220, 350),
        "emp_start": (240, 750), "emp_vac": (520, 750), "emp_wait": (820, 750), "emp_resp": (2920, 750),
        "mod_login": (2060, 1080), "mod_vac": (2340, 1080), "mod_app": (2660, 1080), "mod_dec": (2980, 1080), "end": (3300, 1080),
        "backend_auth": (700, 1500), "backend_vac": (1200, 1500), "backend_app": (1800, 1500), "backend_calc": (2380, 1500), "backend_store": (2920, 1500),
    }
    event(*nodes["start"], "Начало")
    task(*nodes["reg"], "Регистрация аккаунта")
    task(*nodes["login"], "Аутентификация")
    task(*nodes["list"], "Просмотр списка вакансий")
    task(*nodes["detail"], "Открытие карточки вакансии")
    task(*nodes["cart"], "Добавление вакансии в черновик")
    task(*nodes["profile"], "Заполнение профиля")
    task(*nodes["draft"], "Редактирование заявки")
    gateway(*nodes["valid"], "Данные полные?")
    task(*nodes["formed"], "Формирование заявки")
    task(*nodes["watch"], "Просмотр статуса и расчета")
    event(*nodes["emp_start"], "Старт работодателя")
    task(*nodes["emp_vac"], "Создать вакансию и условия доступности")
    task(*nodes["emp_wait"], "Ожидать модерации вакансии")
    task(*nodes["emp_resp"], "Просмотреть отклики")
    task(*nodes["mod_login"], "Войти как модератор")
    task(*nodes["mod_vac"], "Проверить вакансии")
    task(*nodes["mod_app"], "Проверить сформированные заявки")
    gateway(*nodes["mod_dec"], "Решение?")
    event(*nodes["end"], "Процесс завершен", end=True)
    task(*nodes["backend_auth"], "AuthDomain: session-cookie, роли")
    task(*nodes["backend_vac"], "VacancyDomain: фильтры и модерация")
    task(*nodes["backend_app"], "ApplicationDomain: черновик и строки")
    task(*nodes["backend_calc"], "Расчет total_salary и estimated_response_date")
    task(*nodes["backend_store"], "PostgreSQL, Redis, MinIO")

    arrow(draw, [(274, 350), (305, 350)])
    for a1, a2 in [("reg", "login"), ("login", "list"), ("list", "detail"), ("detail", "cart"), ("cart", "profile"), ("profile", "draft"), ("draft", "valid"), ("valid", "formed")]:
        x1, y1 = nodes[a1]; x2, y2 = nodes[a2]
        arrow(draw, [(x1 + 125, y1), (x2 - 125 if a2 != "valid" else x2 - 42, y2)], color="#111827", width=3)
    arrow(draw, [(2380, 392), (2380, 515), (2100, 515), (2100, 393)], color="#111827", text="нет", text_pos=(2220, 493))
    arrow(draw, [(2702, 350), (3095, 350)], color="#111827", text="да", text_pos=(2790, 320))
    arrow(draw, [(274, 750), (395, 750)])
    arrow(draw, [(645, 750), (695, 750)])
    arrow(draw, [(945, 750), (2795, 750)], color="#111827", text="после одобрения", text_pos=(1740, 720))
    arrow(draw, [(2185, 1080), (2215, 1080)])
    arrow(draw, [(2465, 1080), (2535, 1080)])
    arrow(draw, [(2785, 1080), (2938, 1080)])
    arrow(draw, [(3022, 1080), (3266, 1080)], text="finish", text_pos=(3100, 1048))
    arrow(draw, [(2980, 1122), (2980, 1220), (2660, 1220), (2660, 1123)], text="reject", text_pos=(2820, 1190))
    arrow(draw, [(700, 393), (700, 1457)], color="#2563EB", text="POST /login", text_pos=(720, 890))
    arrow(draw, [(980, 393), (1200, 1457)], color="#2563EB", text="GET /vacancies", text_pos=(1005, 970))
    arrow(draw, [(1540, 393), (1800, 1457)], color="#2563EB", text="POST /application-lines", text_pos=(1580, 980))
    arrow(draw, [(2660, 393), (2380, 1457)], color="#2563EB", text="PUT /form", text_pos=(2485, 940))
    arrow(draw, [(2340, 1123), (1200, 1543)], color="#2563EB", text="PUT /vacancies/{id}/moderate", text_pos=(1550, 1300))
    arrow(draw, [(2660, 1123), (1800, 1543)], color="#2563EB", text="GET/PUT /applications", text_pos=(2060, 1305))
    arrow(draw, [(2380, 1543), (2920, 1543)], color="#64748B", text="save metrics/data", text_pos=(2540, 1510))
    img.save(path)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 12) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def add_run(paragraph, text: str, bold: bool = False, italic: bool = False):
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    run.bold = bold
    run.italic = italic
    return run


def add_paragraph(doc: Document, text: str = "", style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    elif style is None:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, text)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    if level == 1:
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        p.paragraph_format.first_line_indent = Cm(0)
        run = add_run(p, text.upper(), bold=True)
        run.font.size = Pt(16)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = add_run(p, text, bold=True)
        run.font.size = Pt(14)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(12)
    add_run(p, text)


def add_picture(doc: Document, path: Path, caption: str, width_cm: float = 15.8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_formula(doc: Document, formula: str, number: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(14)
    table.columns[1].width = Cm(2)
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tc_pr.append(border)
    left, right = table.rows[0].cells
    left.text = ""
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(formula)
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    r.font.size = Pt(14)
    right.text = ""
    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.first_line_indent = Cm(0)
    add_run(p2, number)


def add_table(
    doc: Document,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    font_size: int = 12,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=font_size)
        shade_cell(hdr[i], "EAF2F8")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True
    add_page_number(section)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(14)


def add_landscape_section(doc: Document):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_page_number(section)
    return section


def add_current_page_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.first_line_indent = Cm(0)
    run = add_run(p, text.upper(), bold=True)
    run.font.size = Pt(16)


def section_texts() -> dict[str, list[str]]:
    return {
        "intro": [
            "Актуальность разработки JobAbility связана с тем, что цифровые сервисы поиска работы должны учитывать не только стандартные требования рынка труда, но и особенности трудоустройства людей с ограниченными возможностями здоровья. По данным Всемирной организации здравоохранения, в мире около 1,3 млрд человек, то есть примерно 16 % населения, живут со значительными ограничениями здоровья. Это означает, что задача доступного трудоустройства относится не к узкой социальной нише, а к массовой и устойчивой потребности современного общества [1].",
            "В Российской Федерации правовая база поддержки занятости инвалидов закреплена Федеральным законом от 24.11.1995 № 181-ФЗ «О социальной защите инвалидов в Российской Федерации». Закон предусматривает квотирование рабочих мест, создание специальных условий труда и обязанность учитывать индивидуальные особенности работника [3]. На практике работодателю требуется не просто разместить вакансию, а показать, какие условия адаптации рабочего места доступны: удаленный формат, гибкий график, отсутствие телефонных звонков, доступность офиса, поддержка ассистивных технологий. Соискателю, в свою очередь, важны понятные фильтры, безопасная подача отклика и возможность собрать заявку без лишнего контакта с разными подразделениями.",
            "Цель работы — разработать и развернуть веб-систему JobAbility, предназначенную для публикации доступных вакансий, формирования заявок соискателями, модерации контента и автоматизированного обновления приложения через CI/CD в Kubernetes. Назначение системы состоит в том, чтобы связать соискателей с ограниченными возможностями здоровья, работодателей и модератора в едином информационном контуре.",
            "К нефункциональным требованиям отнесены: доступность интерфейса на настольных и мобильных устройствах; ролевая модель доступа; сохранность данных в PostgreSQL и MinIO; воспроизводимое развертывание в контейнерах; автоматическая проверка и доставка изменений через GitLab CI/CD; наблюдаемость через Prometheus и Grafana; приемлемое время ответа пользовательских операций при демонстрационной нагрузке до 10 HTTP-запросов в секунду.",
            "Для достижения цели решались следующие задачи: исследовать предметную область трудоустройства людей с ОВЗ; определить роли пользователей и бизнес-процесс; спроектировать структуру данных и API; реализовать backend на Django REST Framework; реализовать frontend на React; подготовить Docker-образы; описать Kubernetes-манифесты для приложения и инфраструктурных сервисов; настроить GitLab Runner и pipeline build-test-upload-deploy; проверить работу приложения, мониторинга и хранения медиафайлов.",
        ],
        "domain": [
            "Предметная область рассматривается на примере условного подразделения «Центр инклюзивного найма» в компании JobAbility Partners. В подразделении участвуют HR-менеджер по инклюзивному найму, специалист по адаптации рабочих мест, модератор платформы и руководитель отдела подбора. HR-менеджер публикует вакансии, специалист по адаптации указывает условия доступности рабочего места, модератор проверяет корректность описаний, а руководитель анализирует отклики и эффективность подбора. Такой пример выбран потому, что он отражает реальную организационную цепочку: вакансия не может быть опубликована без проверки, а заявка соискателя должна пройти контролируемый статусный цикл.",
            "Система JobAbility решает прикладную задачу: уменьшить разрыв между формальным размещением вакансии и фактической пригодностью рабочего места для человека с конкретными ограничениями здоровья. В обычных сервисах поиска работы поле «условия труда» часто описано произвольно. В JobAbility оно выделяется как самостоятельная характеристика вакансии: график, город, зарплата, поддержка инвалидности, описание доступности, изображение или видео. Это позволяет пользователю быстрее определить релевантность вакансии.",
            "Ключевой расчет в предметной области — расчет суммарной ожидаемой зарплаты по заявке. Если заявка содержит n выбранных вакансий, а по каждой строке указано количество позиций q_i и зарплата вакансии s_i, то итоговое значение определяется формулой (1):",
            "S = Σ(q_i × s_i), i = 1..n. (1)",
            "В системе эта формула реализована в backend-слое: при редактировании строк заявки пересчитывается line_salary_total, а при формировании заявки пересчитывается total_salary. Дополнительно рассчитывается ожидаемая дата ответа: дата формирования заявки плюс min(30, max(1, n) × 3) дней. Такой расчет не является юридическим сроком, но дает пользователю ориентир по времени обработки.",
            "Бизнес-процесс начинается с публикации вакансии работодателем. Вакансия создается в статусе PENDING, после чего модератор проверяет корректность описания, отсутствие дискриминационных формулировок и соответствие тематике сервиса. После одобрения вакансия становится активной и доступной в публичном каталоге. Соискатель просматривает список вакансий, применяет фильтры, открывает карточку вакансии и добавляет подходящую позицию в черновик заявки. Если черновика нет, backend создает его автоматически.",
            "Затем соискатель заполняет профиль: ФИО, телефон, город, возраст, пол и категорию инвалидности. Эти данные нужны не для публичного отображения, а для корректного оформления заявки. После выбора вакансий пользователь редактирует строки заявки: количество, комментарий, признак основной позиции и порядок. Когда все данные заполнены, заявка переводится из DRAFT в FORMED. На этом этапе система фиксирует дату формирования, рассчитывает сумму и ожидаемую дату ответа.",
            "Модератор видит сформированные заявки во втором фронтенд-сценарии — кабинете модератора. Он открывает заявку, проверяет состав вакансий и может завершить ее со статусом FINISHED либо отклонить со статусом REJECTED с пояснением. Работодатель в своем кабинете видит отклики на созданные им вакансии и может использовать их как основу для дальнейшего контакта с соискателем. Таким образом, сервис отделяет публичный каталог вакансий, личные данные пользователя и административную обработку заявок.",
            "Описание процесса показывает, что JobAbility не является простой витриной вакансий. Это ролевая информационная система: гость получает доступ к каталогу, соискатель формирует заявку, работодатель управляет своими вакансиями, модератор отвечает за качество данных и завершение процесса. Наличие статусов и модерации снижает риск появления некорректных вакансий, а хранение медиафайлов в MinIO позволяет дополнять карточки визуальными материалами без нагрузки на базу данных.",
        ],
        "architecture": [
            "Архитектура JobAbility построена как контейнеризованная веб-система. Пользователь взаимодействует с React SPA, собранным Vite и размещенным в контейнере Nginx. Статический web-сервер обслуживает HTML, CSS и JavaScript, а запросы к /api, /admin, /swagger и /metrics передаются в backend-сервис. Backend реализован на Django 5.2 и Django REST Framework. Он отвечает за бизнес-логику, авторизацию через cookie-session, валидацию входных данных, работу с вакансиями, заявками и профилями.",
            "В Kubernetes приложение размещается в namespace jobability. Внутри кластера работают pods frontend, backend, postgres, redis, minio, prometheus, grafana и adminer. Traefik Ingress принимает HTTP-трафик на 80 порту и направляет его в frontend service. Backend обращается к PostgreSQL по TCP 5432, к Redis по TCP 6379, к MinIO по HTTP/S3 API на 9000 порту. Prometheus периодически опрашивает endpoint /metrics, а Grafana использует Prometheus как datasource.",
            "CI/CD реализован в GitLab. Pipeline выполняется только для основной ветки и состоит из стадий build, test, upload и deploy. На стадии build собираются Docker-образы backend и frontend, tagged коротким SHA коммита. На стадии test backend-образ запускает Django-тесты. На стадии upload образы импортируются в containerd хранилище k3s. На стадии deploy применяются манифесты Kubernetes, выполняется bootstrap MinIO, обновляются образы deployment/backend и deployment/frontend, затем проверяется rollout status.",
            "Расчет аппаратных требований выполнен для демонстрационного стенда на 1000 зарегистрированных пользователей, пиковую активность 5 % и среднюю интенсивность 0,2 запроса в секунду на активного пользователя. Пиковая нагрузка составляет RPS = 1000 × 0,05 × 0,2 = 10 запросов/с. При расчетной стоимости backend-запроса 25 мс CPU и целевой утилизации 65 % требуется ceil(10 × 0,025 / 0,65) = 1 vCPU для приложения. С учетом PostgreSQL, Redis, MinIO, мониторинга, k3s и резервирования требуется не менее 2 vCPU.",
            "По памяти минимальные requests Kubernetes составляют около 608 MiB для прикладных и инфраструктурных pods, но реальные limits дают около 2,4 GiB. С учетом k3s, containerd, системных служб Ubuntu, page cache и 30 % резерва рекомендуется 6-8 GiB RAM. Использованная виртуальная машина с 7,1 GiB RAM и диском 23 GiB удовлетворяет демонстрационному контуру.",
            "Прирост базы данных рассчитывается отдельно от медиафайлов, так как фотографии и видео хранятся в MinIO. Если один пользователь занимает около 2 KiB, одна вакансия около 4 KiB, одна заявка около 3 KiB, а одна строка заявки около 1 KiB, то при 1000 пользователях, 500 вакансиях, 1500 заявках и среднем числе 3 строки на заявку годовой объем PostgreSQL составит примерно 1000×2 + 500×4 + 1500×3 + 4500×1 = 13000 KiB до индексов и служебных данных. С коэффициентом 3 на индексы, TOAST и служебные таблицы достаточно 40-60 MiB. Основной рост ожидается в MinIO: при 500 изображениях по 300 KiB потребуется около 150 MiB, поэтому PVC MinIO выделен отдельно.",
        ],
        "algorithms": [
            "Основной алгоритм пользовательского сценария представлен последовательностью HTTP-запросов между браузером, frontend-страницами, API-клиентами и доменными методами backend. Сначала пользователь проходит аутентификацию: frontend отправляет POST /api/users/login/ с username и password, backend проверяет учетные данные через Django authentication и возвращает сведения о текущем пользователе. Затем выполняется GET /api/users/me/, чтобы восстановить состояние сессии после обновления страницы.",
            "После входа соискатель получает список вакансий через GET /api/vacancies/. В запрос могут передаваться search, min_price, max_price, date_from и date_to. Backend строит queryset только по активным и одобренным вакансиям, применяет фильтры и возвращает массив карточек с image_url, schedule и disability_support. Параллельно интерфейс запрашивает состояние черновой заявки через GET /api/applications/cart/. Если черновика нет, возвращается пустая иконка корзины: application_id = null и items_count = 0.",
            "При добавлении вакансии frontend вызывает POST /api/application-lines/ с vacancy_id и qty. Backend проверяет роль applicant, создает черновую заявку при необходимости, добавляет строку ApplicationVacancy и возвращает идентификатор заявки и строку. Далее пользователь открывает черновик через GET /api/applications/{id}/, редактирует контактные данные и сопроводительное письмо через PUT /api/applications/{id}/, изменяет строки через PUT /api/application-lines/ и при необходимости удаляет их через DELETE /api/application-lines/.",
            "Формирование заявки выполняется запросом PUT /api/applications/{id}/form/. Внутри backend проверяет статус DRAFT, наличие хотя бы одной строки, заполненность ФИО, телефона и города. После успешной проверки строки получают line_salary_total, заявка получает total_salary, formed_at и estimated_response_date, а статус меняется на FORMED. Пользовательский список заявок загружается через GET /api/applications/ и может фильтроваться по status, date_from и date_to.",
            "Во втором фронтенде — кабинете модератора — выполняется вход модератора, затем GET /api/applications/?status=FORMED. Модератор открывает заявку и отправляет PUT /api/applications/{id}/moderate/ с action=finish или action=reject и moderator_note. Backend проверяет роль moderator и допустимый переход состояния. После завершения соискатель при следующем GET /api/applications/ получает список заявок с рассчитанными значениями total_sum, lines_count и calculated_lines_count.",
        ],
        "interface": [
            "Интерфейс JobAbility реализован как одностраничное React-приложение с маршрутизацией react-router-dom. Главная страница объясняет назначение системы: трудоустройство людей с ограниченными возможностями, поиск доступных вакансий и подача откликов. На ней пользователь переходит к каталогу вакансий, регистрации или входу.",
            "Страница регистрации позволяет создать учетную запись соискателя или работодателя. Пользователь вводит имя, фамилию, username, email, пароль и выбирает тип аккаунта. После успешной регистрации система просит выполнить вход вручную, что снижает риск случайного открытия защищенного кабинета без осознанной авторизации.",
            "Страница входа принимает логин и пароль. После успешной авторизации frontend получает роль пользователя и перенаправляет его в соответствующий кабинет. Для соискателя доступен кабинет соискателя, заявки и карточки вакансий; для работодателя — создание вакансий и просмотр откликов; для модератора — проверка вакансий и заявок.",
            "Каталог вакансий содержит карточки с названием, компанией, городом, зарплатой, графиком и условиями доступности. Пользователь может искать вакансии по тексту, ограничивать диапазон зарплаты и даты публикации. На мобильном устройстве карточки и фильтры перестраиваются в вертикальный формат.",
            "Страница детальной вакансии показывает полное описание, визуальные материалы, условия адаптации и кнопку добавления вакансии в заявку. Если пользователь не авторизован, интерфейс предлагает войти. Если пользователь является соискателем, кнопка добавляет позицию в черновик заявки.",
            "Кабинет соискателя показывает профиль, последние заявки и переход к черновику. В профиле редактируются ФИО, телефон, город, возраст, пол и категория инвалидности. Страница списка заявок позволяет отфильтровать заявки по статусу и датам, а детальная страница заявки дает возможность редактировать строки, комментарии, контактный email и сопроводительное письмо.",
            "Кабинет работодателя предназначен для публикации вакансий. Работодатель вводит название, компанию, город, зарплату, график, условия доступности, описание и может приложить изображение или видео. Новая вакансия отправляется на модерацию и не попадает в публичный каталог до одобрения. Отдельная страница откликов работодателя показывает заявки, в которых присутствуют вакансии данного работодателя.",
            "Кабинет модератора объединяет две административные задачи: модерацию вакансий и обработку сформированных заявок. Модератор одобряет или отклоняет вакансии, добавляет комментарий, открывает сформированные заявки и завершает их после проверки. Интерфейс администратора Django и Swagger доступны как вспомогательные окна для демонстрации и проверки API.",
        ],
        "conclusion": [
            "В ходе работы была разработана и развернута система JobAbility — платформа трудоустройства для людей с ограниченными возможностями. Реализованы роли гостя, соискателя, работодателя и модератора; создан каталог вакансий; реализованы заявки со статусами DRAFT, FORMED, FINISHED, REJECTED и DELETED; добавлена модерация вакансий и заявок; настроено хранение медиафайлов в MinIO и кэширование/сессии через Redis.",
            "Backend реализован на Django REST Framework, frontend — на React/Vite. Для эксплуатации подготовлены Dockerfile, Kubernetes-манифесты, ConfigMap, Secret, Service, Deployment, Ingress и PVC. CI/CD в GitLab автоматически собирает образы, запускает тесты, импортирует образы в k3s и обновляет deployments. Дополнительно развернуты Prometheus, Grafana и Adminer.",
            "Результатом является работоспособная система, доступная в локальном контуре виртуальной машины и временно публикуемая через Cloudflare Tunnel. Статическая демонстрационная версия frontend размещается на GitHub Pages: https://xflame0xx.github.io/JobAbility/. Репозитории проекта: GitHub frontend/исходники — https://github.com/xflame0xx/JobAbility, GitLab CI/CD — https://gitlab.com/xflame0xx1/jobability.",
        ],
    }


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)

    use_case = IMG_DIR / "use_case.png"
    state = IMG_DIR / "state.png"
    er = IMG_DIR / "er.png"
    bpmn = IMG_DIR / "bpmn.png"
    make_use_case(use_case)
    make_state(state)
    make_er(er)
    make_bpmn(bpmn)

    class_img = ROOT / "docs" / "diagram_images" / "jobability_class_diagram.png"
    deploy_img = ROOT / "docs" / "diagram_images" / "jobability_deployment_diagram.png"
    seq_img = ROOT / "docs" / "diagram_images" / "jobability_sequence_diagram.png"
    seq_part_1 = ROOT / "docs" / "diagram_images" / "jobability_sequence_diagram_part_1.png"
    seq_part_2 = ROOT / "docs" / "diagram_images" / "jobability_sequence_diagram_part_2.png"

    doc = Document()
    configure_doc(doc)

    # Title page
    for text, size, bold in [
        ("Министерство науки и высшего образования Российской Федерации", 14, False),
        ("МГТУ им. Н. Э. Баумана", 14, False),
        ("Кафедра ИУ-5 «Системы обработки информации и управления»", 14, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = add_run(p, text, bold=bold)
        r.font.size = Pt(size)
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = add_run(p, "РАСЧЕТНО-ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", bold=True)
    r.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = add_run(p, "к итоговому проекту", bold=True)
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    add_run(p, "Тема: «JobAbility — система трудоустройства для людей с ограниченными возможностями»", bold=True)
    for _ in range(8):
        doc.add_paragraph()
    for line in [
        "Студент: ________________________________",
        "Группа: _________________________________",
        "Руководитель: ___________________________",
        "Москва, 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if "Москва" not in line else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        add_run(p, line)

    add_heading(doc, "Аннотация")
    for text in [
        "В расчетно-пояснительной записке описана разработка информационной системы JobAbility, предназначенной для поиска доступных вакансий, формирования заявок соискателями, публикации вакансий работодателями и модерации данных. Рассмотрены предметная область, бизнес-процесс, архитектура web-приложения, структура данных, методы web-сервиса, алгоритм HTTP-взаимодействия и пользовательский интерфейс.",
        "Система реализована как React SPA и Django REST API, развернутые в Kubernetes на виртуальной машине. Для хранения данных используется PostgreSQL, для сессий и кэша — Redis, для медиафайлов — MinIO. Доставка изменений выполняется через GitLab CI/CD с автоматической сборкой Docker-образов, запуском тестов и обновлением Kubernetes deployments.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "Содержание")
    for i, item in enumerate([
        "Введение",
        "Предметная область",
        "Архитектура",
        "Алгоритмы",
        "Описание интерфейса",
        "Заключение",
        "Список использованных источников",
        "Приложение А. Техническое задание",
        "Приложение Б. Методы веб-сервиса",
    ], 1):
        add_paragraph(doc, f"{i}. {item}", align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "Определения, обозначения и сокращения")
    add_table(
        doc,
        "Таблица 1 — Термины и сокращения",
        ["Термин", "Описание"],
        [
            ["ОВЗ", "Ограниченные возможности здоровья."],
            ["SPA", "Single Page Application — одностраничное web-приложение."],
            ["API", "Интерфейс программного взаимодействия между frontend и backend."],
            ["CI/CD", "Непрерывная интеграция и доставка изменений."],
            ["PVC", "PersistentVolumeClaim — запрос постоянного тома в Kubernetes."],
            ["RPS", "Requests per second — количество HTTP-запросов в секунду."],
        ],
    )

    texts = section_texts()

    add_heading(doc, "Введение")
    for t in texts["intro"]:
        add_paragraph(doc, t)

    add_heading(doc, "Предметная область")
    for t in texts["domain"][:5]:
        add_paragraph(doc, t)
    add_picture(doc, use_case, "Рисунок 1 — Диаграмма прецедентов JobAbility", 15.8)
    add_picture(doc, state, "Рисунок 2 — Диаграмма состояний заявки", 15.8)
    for t in texts["domain"][5:]:
        add_paragraph(doc, t)
    add_picture(doc, bpmn, "Рисунок 3 — BPMN бизнес-процесс трудоустройства", 15.8)

    add_heading(doc, "Архитектура")
    for t in texts["architecture"][:3]:
        add_paragraph(doc, t)
    add_picture(doc, deploy_img, "Рисунок 4 — Диаграмма развертывания JobAbility", 15.8)
    add_picture(doc, er, "Рисунок 5 — ER-диаграмма данных", 15.8)
    add_table(
        doc,
        "Таблица 2 — Назначение таблиц базы данных",
        ["Таблица", "Назначение"],
        [
            ["auth_user", "Стандартная таблица Django для учетных записей, паролей, email и ФИО."],
            ["core_useraccount", "Расширение пользователя: роль соискателя, работодателя или модератора."],
            ["core_applicantprofile", "Профиль соискателя: контактные и анкетные данные."],
            ["core_vacancy", "Вакансии, сведения о работодателе, модерации, медиа и условиях доступности."],
            ["core_application", "Заявки соискателей, статусы, даты, контактные данные, расчет total_salary."],
            ["core_applicationvacancy", "Строки заявки: выбранные вакансии, количество, комментарии и расчет по строке."],
        ],
    )
    table_rows = [
        ["auth_user", "id", "bigint", "Первичный ключ пользователя"],
        ["auth_user", "username", "varchar(150)", "Логин"],
        ["auth_user", "password", "varchar(128)", "Хэш пароля"],
        ["auth_user", "email", "varchar(254)", "Электронная почта"],
        ["core_useraccount", "user_id", "bigint FK", "Связь 1:1 с auth_user"],
        ["core_useraccount", "role", "varchar(32)", "Роль: applicant, employer, moderator"],
        ["core_applicantprofile", "full_name", "varchar(255)", "ФИО соискателя"],
        ["core_applicantprofile", "phone, city, age", "varchar/int", "Контактные и анкетные данные"],
        ["core_applicantprofile", "disability_category", "varchar(16)", "Категория инвалидности"],
        ["core_vacancy", "title, company, city", "varchar", "Основные сведения о вакансии"],
        ["core_vacancy", "salary", "integer", "Зарплата для расчетов"],
        ["core_vacancy", "moderation_status", "varchar(16)", "PENDING, APPROVED, REJECTED"],
        ["core_vacancy", "image, video", "file", "Медиафайлы в MinIO"],
        ["core_application", "status", "varchar(16)", "DRAFT, FORMED, FINISHED, REJECTED, DELETED"],
        ["core_application", "formed_at, completed_at", "datetime", "Даты формирования и завершения"],
        ["core_application", "total_salary", "integer", "Итоговый расчет по строкам заявки"],
        ["core_applicationvacancy", "application_id, vacancy_id", "bigint FK", "Связь заявки и вакансии"],
        ["core_applicationvacancy", "qty", "integer", "Количество выбранных позиций"],
        ["core_applicationvacancy", "line_salary_total", "integer", "Расчет по строке заявки"],
    ]
    add_table(doc, "Таблица 3 — Описание основных колонок таблиц", ["Таблица", "Колонка", "Тип", "Описание"], table_rows)
    add_picture(doc, class_img, "Рисунок 6 — Диаграмма классов frontend/backend доменов", 15.8)
    for t in texts["architecture"][3:]:
        add_paragraph(doc, t)

    add_heading(doc, "Алгоритмы")
    for t in texts["algorithms"]:
        add_paragraph(doc, t)
    if seq_part_1.exists() and seq_part_2.exists():
        add_picture(doc, seq_part_1, "Рисунок 7 — Диаграмма последовательности HTTP-запросов (часть 1)", 15.0)
        add_picture(doc, seq_part_2, "Рисунок 8 — Диаграмма последовательности HTTP-запросов (часть 2)", 15.0)
    else:
        add_picture(doc, seq_img, "Рисунок 7 — Диаграмма последовательности HTTP-запросов", 15.0)

    add_heading(doc, "Описание интерфейса")
    for t in texts["interface"]:
        add_paragraph(doc, t)
    add_table(
        doc,
        "Таблица 4 — Перечень основных окон интерфейса",
        ["Окно", "Маршрут", "Пользовательские действия"],
        [
            ["Главная", "/", "Переход к каталогу, входу и регистрации."],
            ["Регистрация", "/register", "Создание аккаунта соискателя или работодателя."],
            ["Вход", "/login", "Авторизация и переход по роли."],
            ["Вакансии", "/vacancies", "Поиск, фильтрация, открытие карточек."],
            ["Карточка вакансии", "/vacancies/:id", "Просмотр описания и добавление в заявку."],
            ["Заявки", "/applications", "Список заявок и фильтры по статусам."],
            ["Детальная заявка", "/applications/:id", "Редактирование, формирование, модерация."],
            ["Кабинет соискателя", "/cabinet/applicant", "Редактирование профиля и переход к заявкам."],
            ["Кабинет работодателя", "/cabinet/employer", "Создание вакансий и просмотр статусов модерации."],
            ["Отклики работодателя", "/cabinet/employer/responses", "Просмотр заявок по вакансиям работодателя."],
            ["Кабинет модератора", "/cabinet/moderator", "Модерация вакансий и заявок."],
            ["Swagger", "/swagger/", "Проверка и демонстрация API."],
            ["Django Admin", "/admin/", "Административная проверка данных."],
        ],
    )

    add_heading(doc, "Заключение")
    for t in texts["conclusion"]:
        add_paragraph(doc, t)

    add_heading(doc, "Список использованных источников")
    sources = [
        "World Health Organization. Disability: fact sheet. URL: https://www.who.int/news-room/fact-sheets/detail/disability-and-health (дата обращения: 29.05.2026).",
        "International Labour Organization. Disability inclusion at work. URL: https://www.ilo.org/topics/disability-and-work (дата обращения: 29.05.2026).",
        "Федеральный закон от 24.11.1995 № 181-ФЗ «О социальной защите инвалидов в Российской Федерации». URL: https://www.consultant.ru/document/cons_doc_LAW_8559/ (дата обращения: 29.05.2026).",
        "Django Software Foundation. Django documentation 5.2. URL: https://docs.djangoproject.com/en/5.2/ (дата обращения: 29.05.2026).",
        "Django REST framework. API Guide. URL: https://www.django-rest-framework.org/api-guide/ (дата обращения: 29.05.2026).",
        "React documentation. URL: https://react.dev/ (дата обращения: 29.05.2026).",
        "Kubernetes Documentation. Concepts: workloads, services, ingress. URL: https://kubernetes.io/docs/concepts/ (дата обращения: 29.05.2026).",
        "Docker Documentation. Build and Dockerfile reference. URL: https://docs.docker.com/build/ (дата обращения: 29.05.2026).",
        "PostgreSQL Documentation. URL: https://www.postgresql.org/docs/ (дата обращения: 29.05.2026).",
        "GitLab Documentation. CI/CD pipelines. URL: https://docs.gitlab.com/ci/ (дата обращения: 29.05.2026).",
        "MinIO Documentation. URL: https://min.io/docs/minio/kubernetes/upstream/ (дата обращения: 29.05.2026).",
        "Redis Documentation. URL: https://redis.io/docs/latest/ (дата обращения: 29.05.2026).",
    ]
    for i, s in enumerate(sources, 1):
        add_paragraph(doc, f"{i}. {s}")

    add_heading(doc, "Приложение А. Техническое задание")
    for t in [
        "Наименование системы: JobAbility — web-система трудоустройства для людей с ограниченными возможностями.",
        "Основание разработки: итоговый проект по дисциплине, предусматривающий контейнеризацию проекта, развертывание в Kubernetes, настройку CI/CD и подготовку отчетной документации.",
        "Назначение: предоставление соискателям доступного каталога вакансий, работодателям — инструмента публикации вакансий, модератору — инструмента контроля качества данных и завершения заявок.",
        "Функциональные требования: регистрация и вход; роли applicant, employer, moderator; просмотр и фильтрация вакансий; создание вакансии работодателем; модерация вакансии; добавление вакансии в заявку; редактирование профиля и черновика; формирование заявки; список заявок; модерация заявок; загрузка изображений и видео; Swagger-документация API; мониторинг метрик.",
        "Требования к надежности: данные PostgreSQL, Redis, MinIO, Prometheus и Grafana должны храниться на PVC; deployments должны иметь readinessProbe; секреты должны храниться в Kubernetes Secret; конфигурация — в ConfigMap.",
        "Требования к развертыванию: система запускается в namespace jobability кластера k3s; frontend обслуживается Nginx; backend запускается через Gunicorn/Django; Ingress направляет внешний HTTP-трафик на frontend; CI/CD выполняет build, test, upload и deploy.",
    ]:
        add_paragraph(doc, t)

    add_heading(doc, "Приложение Б. Методы веб-сервиса")
    api_rows = [
        ["1", "Auth", "GET", "/api/users/me/", "Текущий пользователь", "-", "id:int, username:str, role:str, full_name:str, email:str"],
        ["2", "Auth", "POST", "/api/users/register/", "Регистрация", "username:str, password:str, first_name:str, last_name:str, email:str, role:str", "id:int, username:str, role:str"],
        ["3", "Auth", "POST", "/api/users/login/", "Вход", "username:str, password:str", "id:int, username:str, role:str, session_key:str"],
        ["4", "Auth", "POST", "/api/users/logout/", "Выход", "-", "message:str"],
        ["5", "Auth", "PUT", "/api/users/password/", "Смена пароля", "old_password:str, new_password:str, new_password_repeat:str", "message:str"],
        ["6", "Vacancy", "GET", "/api/vacancies/", "Список вакансий", "search:str, min_price:int, max_price:int, date_from:date, date_to:date", "Vacancy[]"],
        ["7", "Vacancy", "POST", "/api/vacancies/", "Создание вакансии", "title, company, city:str; salary:int; image:file; video:file", "Vacancy"],
        ["8", "Vacancy", "GET", "/api/vacancies/{id}/", "Карточка вакансии", "id:int", "Vacancy"],
        ["9", "Vacancy", "GET", "/api/vacancies/mine/", "Вакансии работодателя", "-", "Vacancy[]"],
        ["10", "Vacancy", "GET", "/api/vacancies/pending/", "Вакансии на модерации", "-", "Vacancy[]"],
        ["11", "Vacancy", "PUT", "/api/vacancies/{id}/moderate/", "Модерация вакансии", "action:str, moderation_note:str", "Vacancy"],
        ["12", "Profile", "GET", "/api/users/profile/", "Профиль соискателя", "-", "full_name, phone, city, age, gender, disability_category"],
        ["13", "Profile", "PUT", "/api/users/profile/", "Обновление профиля", "full_name, phone, city, age, gender, disability_category", "ApplicantProfile"],
        ["14", "Application", "GET", "/api/applications/cart/", "Иконка черновика", "-", "application_id:int|null, items_count:int"],
        ["15", "Application", "POST", "/api/application-lines/", "Добавить вакансию", "vacancy_id:int, qty:int", "application_id:int, line:ApplicationLine"],
        ["16", "Application", "PUT", "/api/application-lines/", "Изменить строку", "vacancy_id:int, qty:int, comment:str, is_main:bool, order_index:int", "ApplicationLine"],
        ["17", "Application", "DELETE", "/api/application-lines/", "Удалить строку", "vacancy_id:int", "204 No Content"],
        ["18", "Application", "GET", "/api/applications/", "Список заявок", "status:str, date_from:date, date_to:date", "ApplicationList[]"],
        ["19", "Application", "GET", "/api/applications/{id}/", "Детальная заявка", "id:int", "ApplicationDetail"],
        ["20", "Application", "PUT", "/api/applications/{id}/", "Редактирование заявки", "profile fields, contact_email:str, cover_letter:str", "ApplicationDetail"],
        ["21", "Application", "PUT", "/api/applications/{id}/form/", "Формирование заявки", "id:int", "ApplicationDetail"],
        ["22", "Application", "PUT", "/api/applications/{id}/moderate/", "Завершение/отклонение", "action:finish|reject, moderator_note:str", "ApplicationDetail"],
        ["23", "Application", "DELETE", "/api/applications/{id}/delete/", "Удаление черновика", "id:int", "204 No Content"],
        ["24", "Employer", "GET", "/api/applications/employer-responses/", "Отклики работодателя", "-", "ApplicationList[]"],
        ["25", "System", "GET", "/api/schema/", "OpenAPI schema", "-", "OpenAPI JSON"],
        ["26", "System", "GET", "/swagger/", "Swagger UI", "-", "HTML"],
        ["27", "System", "GET", "/metrics/", "Метрики Prometheus", "-", "text/plain metrics"],
    ]
    grouped: dict[str, list[list[str]]] = {}
    for row in api_rows:
        grouped.setdefault(row[1], []).append([row[0], row[2], row[3], row[4], row[5], row[6]])
    table_no = 1
    for domain, rows in grouped.items():
        add_paragraph(doc, f"Домен {domain}.", align=WD_ALIGN_PARAGRAPH.LEFT)
        add_table(
            doc,
            f"Таблица Б.{table_no} — Методы домена {domain}",
            ["№", "Метод", "URL", "Описание", "Входные данные", "Выходные данные"],
            rows,
            widths=[0.9, 1.4, 3.6, 3.7, 4.4, 3.8],
            font_size=9,
        )
        table_no += 1

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
